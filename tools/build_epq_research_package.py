#!/usr/bin/env python3
"""Build the revised EPQ dissertation and its episode-one video script."""

from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "assets" / "documents"
OUT.mkdir(parents=True, exist_ok=True)

# Explicit standard_business_brief-inspired document tokens.
INK = RGBColor(15, 23, 42)
MUTED = RGBColor(71, 85, 105)
ACCENT = RGBColor(13, 148, 136)
ACCENT_DARK = RGBColor(15, 118, 110)
PANEL = "F1F5F9"
RULE = "CBD5E1"
WHITE = RGBColor(255, 255, 255)
FONT = "Aptos"
DISPLAY = "Aptos Display"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs[edge]
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_char2])


def add_hyperlink(paragraph, text: str, url: str, color: str = "0F766E"):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    new_run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([run_color, underline])
    new_run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def set_core_properties(doc: Document, title: str, subject: str) -> None:
    props = doc.core_properties
    props.author = "Nathan Brown-Bennett"
    props.last_modified_by = "Nathan Brown-Bennett"
    props.title = title
    props.subject = subject
    props.keywords = "cybersecurity, research, EPQ, risk management, secure by design"
    props.comments = ""
    props.category = "Research"


def configure_document(doc: Document, title: str, subject: str) -> None:
    set_core_properties(doc, title, subject)
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.16
    normal.paragraph_format.widow_control = True

    for style_name, size, color, before, after in (
        ("Title", 30, INK, 0, 12),
        ("Subtitle", 13, MUTED, 0, 18),
        ("Heading 1", 19, INK, 18, 8),
        ("Heading 2", 13, ACCENT_DARK, 12, 5),
        ("Heading 3", 11, INK, 9, 4),
    ):
        style = styles[style_name]
        style.font.name = DISPLAY if style_name in {"Title", "Heading 1"} else FONT
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Eyebrow" not in styles:
        eyebrow = styles.add_style("Eyebrow", WD_STYLE_TYPE.PARAGRAPH)
    else:
        eyebrow = styles["Eyebrow"]
    eyebrow.font.name = FONT
    eyebrow.font.size = Pt(9)
    eyebrow.font.bold = True
    eyebrow.font.color.rgb = ACCENT_DARK
    eyebrow.paragraph_format.space_after = Pt(8)

    if "Source note" not in styles:
        source_style = styles.add_style("Source note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        source_style = styles["Source note"]
    source_style.font.name = FONT
    source_style.font.size = Pt(8.5)
    source_style.font.color.rgb = MUTED
    source_style.paragraph_format.space_after = Pt(5)
    source_style.paragraph_format.line_spacing = 1.05

    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "NATHAN BROWN-BENNETT  /  RESEARCH NOTES"
    header_para.style = styles["Source note"]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.style = styles["Source note"]
    footer_para.add_run("2026 REVISED EDITION")
    footer_para.add_run("\t")
    add_page_number(footer_para)


def add_title_page(
    doc: Document,
    eyebrow: str,
    title: str,
    subtitle: str,
    meta_lines: Iterable[str],
    note: str,
    page_break: bool = True,
) -> None:
    p = doc.add_paragraph(style="Eyebrow")
    p.add_run(eyebrow.upper())
    doc.add_paragraph(title, style="Title")
    doc.add_paragraph(subtitle, style="Subtitle")
    doc.add_paragraph("")
    for line in meta_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        r.font.size = Pt(10)
        r.font.color.rgb = MUTED
    doc.add_paragraph("")
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Cm(15.6)
    cell = table.cell(0, 0)
    cell.width = Cm(15.6)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, PANEL)
    set_cell_border(
        cell,
        left={"val": "single", "sz": "18", "color": "0D9488"},
        top={"val": "nil"},
        right={"val": "nil"},
        bottom={"val": "nil"},
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.left_indent = Pt(8)
    r = p.add_run(note)
    r.font.size = Pt(10)
    r.font.color.rgb = INK
    if page_break:
        doc.add_page_break()


def set_section_columns(section, count: int = 2, space_twips: int = 420) -> None:
    """Set Word newspaper-style columns on a section."""
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(space_twips))
    cols.set(qn("w:equalWidth"), "1")


def add_part_heading(doc: Document, number: str, title: str, summary: str) -> None:
    p = doc.add_paragraph(style="Eyebrow")
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(f"PART {number}")
    heading = doc.add_paragraph(title, style="Heading 1")
    heading.paragraph_format.space_before = Pt(0)
    p2 = doc.add_paragraph(summary)
    p2.paragraph_format.space_after = Pt(10)
    p2.runs[0].font.color.rgb = MUTED


def add_labelled_callout(doc: Document, label: str, body: str) -> None:
    def style_paragraph(paragraph, top: bool, bottom: bool) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), PANEL)
        p_pr.append(shading)
        borders = OxmlElement("w:pBdr")
        for edge, size, color in (
            ("left", "14", "0D9488"),
            ("top", "4", RULE),
            ("right", "4", RULE),
            ("bottom", "4", RULE),
        ):
            if edge == "top" and not top:
                continue
            if edge == "bottom" and not bottom:
                continue
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), size)
            element.set(qn("w:space"), "5")
            element.set(qn("w:color"), color)
            borders.append(element)
        p_pr.append(borders)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Pt(8)
    p.paragraph_format.right_indent = Pt(6)
    style_paragraph(p, top=True, bottom=False)
    r = p.add_run(label.upper())
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = ACCENT_DARK
    p2 = doc.add_paragraph(body)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(7)
    p2.paragraph_format.left_indent = Pt(8)
    p2.paragraph_format.right_indent = Pt(6)
    p2.paragraph_format.line_spacing = 1.1
    style_paragraph(p2, top=False, bottom=True)
    keep_with_next(p)


def add_bullets(doc: Document, items: Iterable[str], level: int = 0) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_evidence_table(doc: Document) -> None:
    data = [
        (
            "43%",
            "UK businesses identifying a breach or attack in the preceding 12 months",
            "DSIT & Home Office, 2026",
        ),
        (
            "204",
            "nationally significant incidents handled by the NCSC in 2024–25",
            "NCSC, 2025",
        ),
        (
            "30%",
            "businesses conducting a cyber-risk assessment",
            "DSIT & Home Office, 2026",
        ),
        (
            "25%",
            "businesses with a formal incident-response plan",
            "DSIT & Home Office, 2026",
        ),
    ]
    for metric, meaning, source in data:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{metric} ")
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = ACCENT_DARK
        p.add_run(f"{meaning}. ")
        s = p.add_run(source)
        s.italic = True
        s.font.color.rgb = MUTED
    doc.add_paragraph(
        "Interpretation: attack prevalence measures exposure and detection, not the absence of defensive value. "
        "The more revealing gap is between recognising cyber risk and consistently governing, testing and recovering from it.",
        style="Source note",
    )


def add_reference(doc: Document, citation: str, title: str, url: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_after = Pt(6)
    p.add_run(citation + " ").bold = False
    add_hyperlink(p, title, url)


def build_dissertation() -> Path:
    doc = Document()
    title = "Can cybersecurity keep pace with modern digital risk?"
    configure_document(
        doc,
        title,
        "A 2026 reappraisal of the author's A-level Extended Project Qualification",
    )
    add_title_page(
        doc,
        "Research Notes / Episode 01",
        title,
        "A 2026 reappraisal of my A-level EPQ: “An investigation into Cyber Security and whether it is secure for modern-day use?”",
        (
            "Nathan Brown-Bennett",
            "Revised research edition · July 2026",
            "Original A-level submission: 2020",
        ),
        "This is a new critical edition, not a retrospective alteration of the assessed A-level submission. "
        "The original is preserved separately. Claims, evidence and references in this edition were re-evaluated in 2026.",
        page_break=False,
    )
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_columns(body_section, 2, 420)
    body_section.top_margin = Cm(1.6)
    body_section.bottom_margin = Cm(1.6)
    body_section.left_margin = Cm(1.7)
    body_section.right_margin = Cm(1.7)
    doc.styles["Normal"].font.size = Pt(9.4)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.07
    doc.styles["Normal"].paragraph_format.space_after = Pt(5)
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(13)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(5)
    doc.styles["Heading 2"].font.size = Pt(11.5)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(8)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(3)

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "My original EPQ asked whether cybersecurity itself was “secure” enough for modern use. "
        "That wording captured a genuine concern but treated security as a binary property: either a system is safe or it has failed. "
        "This revised paper reframes the question as whether contemporary cybersecurity can reduce, govern and recover from digital risk at the pace at which that risk changes. "
        "It revisits the original themes—definitions, public understanding, human behaviour, automation, regulation and education—using evidence available in 2026."
    )
    doc.add_paragraph(
        "The evidence shows a difficult but more useful answer. Threat volume and consequence remain high: 43% of UK businesses identified a breach or attack in the 2025/2026 survey, while the National Cyber Security Centre handled 204 nationally significant incidents during 2024–25. "
        "Yet successful defence is often invisible, and the existence of incidents does not demonstrate that cybersecurity is ineffective. "
        "Effectiveness is better judged through risk ownership, secure-by-design engineering, usable defaults, detection, response and recovery. "
        "The central conclusion is that cybersecurity can keep pace only when it is treated as a continuous socio-technical capability rather than a product, a one-off control, or a demand that individual users never make mistakes."
    )
    doc.add_heading("Research integrity and provenance", level=1)
    doc.add_paragraph(
        "The historical source is the 22-page Word document submitted for my A-level Extended Project Qualification. "
        "Two copies supplied for this revision were byte-for-byte identical. A PDF supplied alongside them contained only one rendered page and is therefore not a complete archival copy. "
        "The portfolio previously linked to a different planning document about PC security and practicality; this edition corrects that mismatch."
    )
    add_labelled_callout(
        doc,
        "What has and has not been changed",
        "This edition preserves the original intellectual starting point but rewrites the argument, replaces weak or stale web references with authoritative and current material, and adds explicit limitations. "
        "No new participant study has been conducted. The original small Google Forms exercise is discussed only as historical evidence of how the project was developed; it is not presented as representative research.",
    )

    doc.add_heading("Report structure", level=1)
    add_numbered(
        doc,
        (
            "Part I — Define the field and repair the research question",
            "Part II — Test the argument against current evidence",
            "Part III — Explain what effective practice requires",
            "Part IV — Trace the research into a product direction",
        ),
    )

    add_part_heading(
        doc,
        "I",
        "Define the field",
        "The terminology matters because each field protects a different boundary and answers a different operational question.",
    )
    doc.add_heading("1. What security means", level=1)
    doc.add_paragraph(
        "Security is the managed reduction of unacceptable harm. It is not the promise that nothing bad can happen. "
        "A security decision identifies something of value, the threats and weaknesses that could affect it, the likely consequence, and the controls proportionate to that risk."
    )
    doc.add_paragraph(
        "Four properties recur across the security fields: confidentiality—preventing unauthorised disclosure; integrity—preventing unauthorised or accidental change; availability—keeping authorised access and essential service; and resilience—containing disruption and restoring trusted operation. "
        "Authenticity, accountability, safety and privacy may also matter depending on the system."
    )
    add_labelled_callout(
        doc,
        "A practical test",
        "What are we protecting? From which harm? Who owns the decision? Which control reduces the risk? How will we know it works? What happens when it fails?",
    )
    doc.add_heading("2. Why the original question needed reframing", level=1)
    doc.add_paragraph(
        "The phrase “whether cybersecurity is secure” sounds intuitive, but it combines two different ideas. "
        "Cybersecurity is the collection of people, processes, technologies and decisions used to manage digital risk. "
        "A particular system can be more or less resistant to known threats; the discipline used to protect it cannot meaningfully be certified as permanently “secure”. "
        "The original question therefore risked judging an adaptive practice against an impossible condition."
    )
    doc.add_paragraph(
        "The original paper used the history of locks to contrast “perfect security” with modern defensive practice. "
        "That analogy remains useful if its lesson is stated carefully: a lock is not worthless because it can eventually be picked, and a cyber control is not ineffective because a sufficiently capable adversary might bypass it. "
        "Security is contextual. It depends on what is being protected, from whom, for how long, at what cost and with what consequence if the control fails."
    )
    doc.add_paragraph(
        "This is consistent with modern frameworks. NIST’s Cybersecurity Framework 2.0 describes outcomes for governing, identifying, protecting, detecting, responding and recovering from cyber risk (Pascoe, Quinn and Scarfone, 2024). "
        "The addition of “Govern” is especially important. Security is not simply an IT team’s defensive activity; it is a management decision about priorities, responsibilities, dependencies and acceptable risk."
    )
    add_labelled_callout(
        doc,
        "Revised research question",
        "Can contemporary cybersecurity reduce and govern digital risk—and help organisations recover—at the pace at which technology, dependency and adversarial capability change?",
    )

    doc.add_heading("3. What each security field handles", level=1)
    doc.add_paragraph(
        "These terms overlap, but they are not synonyms. A useful distinction begins with the object being protected and the boundary within which the work is performed."
    )
    distinctions = [
        (
            "Computer security",
            "Protects an individual computing system—its operating system, software, memory, storage and local data. Typical controls include patching, secure boot, malware protection, access control and hardening. Example: preventing unauthorised code from gaining administrator access on a laptop.",
        ),
        (
            "IT security",
            "Protects an organisation’s technology estate and the services it operates: devices, servers, networks, applications, identities, cloud platforms and backups. Typical controls include asset management, network segmentation, identity administration, monitoring and recovery. Example: ensuring staff accounts and managed devices can use a business system safely.",
        ),
        (
            "Information security",
            "Protects information in any form—digital, paper or spoken—against loss of confidentiality, integrity or availability. It is broader than technology. Example: controlling who may view a printed personnel file as well as its digital copy.",
        ),
        (
            "Cybersecurity",
            "Manages harm arising through connected digital systems. It includes technology and information, but can also protect people, services, organisations and physical processes reachable through cyberspace. Example: preventing a compromised connected device from disrupting a service or exposing its users.",
        ),
    ]
    for field, boundary in distinctions:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{field}. ")
        r.bold = True
        r.font.color.rgb = ACCENT_DARK
        p.add_run(boundary)
    doc.add_paragraph(
        "In practice, one incident can cross all four boundaries. A vulnerable laptop is a computer-security problem; weak enterprise account management is an IT-security problem; exposed records are an information-security problem; and the connected attack path and resulting harm are a cybersecurity problem.",
        style="Source note",
    )
    doc.add_heading("3.1 Typical roles and deliverables", level=2)
    role_rows = [
        (
            "Computer security",
            "Endpoint or platform security engineer; operating-system specialist; malware analyst. Outputs: hardened builds, patch baselines, endpoint controls and technical findings.",
        ),
        (
            "IT security",
            "Security analyst; identity and access specialist; network or cloud security engineer; security administrator. Outputs: managed access, architecture, monitoring, backup and operational control.",
        ),
        (
            "Information security",
            "Information-security manager; risk and compliance specialist; security governance lead; information asset owner. Outputs: classification, risk treatment, policy, assurance and audit evidence.",
        ),
        (
            "Cybersecurity",
            "SOC analyst; incident responder; threat-intelligence analyst; penetration tester; application-security engineer; security architect. Outputs: threat-led defence, detection, incident containment, testing and recovery improvement.",
        ),
    ]
    for field, roles in role_rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{field}. ")
        r.bold = True
        r.font.color.rgb = ACCENT_DARK
        p.add_run(roles)
    doc.add_paragraph(
        "Job titles are not standardised. A small organisation may combine all of these duties in one IT role; a large organisation may split one duty across several teams. Scope, authority and evidence are more reliable than the title alone.",
        style="Source note",
    )
    doc.add_heading("3.2 Where the fields overlap", level=2)
    add_bullets(
        doc,
        (
            "Identity and access control crosses computer, IT, information and cyber security.",
            "Patching begins as computer or platform security, becomes an IT deployment process, and reduces a cyber attack path.",
            "Logging is configured by technical teams, monitored by security operations, and retained as information and audit evidence.",
            "Incident response joins technical containment, business continuity, information assessment, legal duties and external communication.",
            "Secure product development combines computer security, application security, information protection and threat-led cybersecurity.",
        ),
    )
    doc.add_heading("3.3 One incident, four perspectives", level=2)
    doc.add_paragraph(
        "Consider a staff laptop compromised through a stolen cloud credential. Computer security asks whether the endpoint was patched, hardened and able to isolate malicious activity. "
        "IT security asks how the identity, device, network and cloud service were administered and monitored. "
        "Information security identifies which records were exposed, their sensitivity and the organisation’s obligations. "
        "Cybersecurity reconstructs the connected attack path, contains the adversary, assesses wider harm and improves detection and resilience."
    )
    doc.add_paragraph(
        "The fields therefore overlap by design. Clear handoffs prevent gaps: the endpoint team supplies evidence; IT revokes access and restores service; information owners assess impact; incident responders coordinate containment and lessons learned; leadership accepts or funds the remaining risk."
    )
    doc.add_heading("3.4 The academic boundary", level=2)
    doc.add_paragraph(
        "Von Solms and van Niekerk argue that cybersecurity extends beyond traditional information assets to people and other assets reachable through cyberspace (2013). "
        "That distinction is useful, but organisational teams still share controls and responsibilities. The labels should clarify ownership—not create silos."
    )
    doc.add_heading("3.5 Resilience changes the test", level=2)
    doc.add_paragraph(
        "A prevention-only definition creates a misleading pass/fail test. Modern services depend on cloud platforms, identity providers, software libraries, suppliers, connected devices and human decisions. "
        "Some failures will occur even in mature environments. Resilience asks the harder questions: Was the most likely harm reduced? Was abnormal activity detected? Could the organisation isolate the problem? Were essential services maintained? Could trustworthy operations be restored and lessons applied?"
    )
    doc.add_paragraph(
        "The NCSC’s 2025 Annual Review makes this distinction directly. The agency reported that attackers were improving their ability to cause impact, but also observed that many more attacks fail than succeed and that prepared organisations are better able to continue operating when an attack gets through (NCSC, 2025a). "
        "This means incident counts should inform urgency, not serve as a simple score of whether cybersecurity “works”."
    )

    doc.add_heading("4. What the original project got right", level=1)
    doc.add_paragraph(
        "The A-level paper was written with limited access to academic databases and before several now-central UK policy changes. "
        "Even so, four instincts have survived scrutiny."
    )
    add_bullets(
        doc,
        (
            "Perfect security is not a useful operational goal. Defenders need continuous learning, layered controls and recovery.",
            "Public understanding matters. Security decisions fail when advice is incomprehensible, unaffordable or detached from how people actually work.",
            "Automation is necessary at scale. Machines can process telemetry and repeat controls more consistently than manual work alone.",
            "Cybersecurity connects technology with governance, education and society. It cannot be reduced to antivirus software or a specialist team.",
        ),
    )
    doc.add_paragraph(
        "The original participant exercise also revealed a worthwhile design question: people can feel confident about common cyber terms while interpreting them differently. "
        "However, the study’s sample was small and socially concentrated, the document reports both 20 and 25 responses in different places, and one question assumed a universal minimum password length that does not exist. "
        "The finding is therefore a prompt for further research, not evidence about the general population."
    )

    add_part_heading(
        doc,
        "II",
        "Test the argument",
        "Current evidence shows high exposure, rising consequence and a persistent gap between available practice and consistent adoption.",
    )
    doc.add_heading("5. What has changed since 2020", level=1)
    doc.add_heading("5.1 Digital dependency became operational dependency", level=2)
    doc.add_paragraph(
        "The years following the original submission made digital dependency more visible. Remote work, cloud services, software-as-a-service, connected devices and complex supply chains increased the number of systems that organisations must trust. "
        "Cyber harm is no longer confined to lost files: incidents can interrupt retail, manufacturing, healthcare, education, logistics and public services."
    )
    doc.add_paragraph(
        "ENISA analysed 4,875 incidents between July 2024 and June 2025 for its 2025 Threat Landscape, illustrating both the scale and diversity of the European threat environment (ENISA, 2025). "
        "In the UK, the NCSC handled 429 incidents in 2024–25, of which 204 were nationally significant and 18 highly significant. "
        "The number of nationally significant incidents was more than double the previous year’s 89, although reporting and categorisation data must not be treated as a complete census (NCSC, 2025b)."
    )
    doc.add_heading("5.2 Security expectations moved towards products and leadership", level=2)
    doc.add_paragraph(
        "Responsibility has begun to move upstream. The UK’s consumer connectable product security regime came into effect on 29 April 2024, placing minimum security obligations on manufacturers and supply-chain businesses for relevant smart products (DSIT, 2024). "
        "The UK Cyber Governance Code of Practice, launched in 2025, similarly frames cyber risk as a board responsibility tied to operational and financial viability (DSIT and NCSC, 2025)."
    )
    doc.add_paragraph(
        "For software, secure-by-design guidance now asks vendors to threat-model early, enforce strong authentication for privileged users, remove default passwords, validate input, protect credentials and plan for secure updates (NCSC, 2025c). "
        "This is a decisive shift from telling end users to compensate for unsafe products."
    )

    doc.add_heading("6. What current evidence says about effectiveness", level=1)
    add_evidence_table(doc)
    doc.add_heading("6.1 Exposure remains widespread", level=2)
    doc.add_paragraph(
        "The Cyber Security Breaches Survey 2025/2026 estimates that 43% of UK businesses and 28% of charities identified a breach or attack in the previous 12 months. "
        "Medium and large businesses reported higher prevalence than micro businesses, partly because they present larger targets and often have stronger detection capability (DSIT and Home Office, 2026). "
        "The survey explicitly warns that unidentified or undisclosed incidents mean the true prevalence may be higher."
    )
    doc.add_paragraph(
        "Phishing remained the dominant cyber-crime category among organisations reporting cyber crime. "
        "Yet the same survey found that most organisations affected by their most disruptive breach restored operations within 24 hours. "
        "Those facts can coexist: exposure can be frequent while many incidents are contained. A mature assessment must measure severity, duration, recovery and avoided harm—not just whether an attempt occurred."
    )
    doc.add_heading("6.2 Governance and preparedness are inconsistent", level=2)
    doc.add_paragraph(
        "The more concerning figures describe preparation. Only 30% of businesses reported conducting a cyber-risk assessment, 31% assigned explicit board responsibility, and 25% had a formal incident-response plan. "
        "Forty-five per cent of businesses reported none of the listed incident-response measures. "
        "Even where plans existed, qualitative interviews suggested that some had never been tested (DSIT and Home Office, 2026)."
    )
    doc.add_paragraph(
        "This gap supports a revised conclusion: the limiting factor is not that cybersecurity as a discipline has no effective practices. "
        "It is that proven practices are unevenly adopted, poorly integrated with business decisions, or left untested until a crisis."
    )

    add_part_heading(
        doc,
        "III",
        "Explain effective practice",
        "Security works when products, organisations and people share responsibility through usable controls, evidence and tested recovery.",
    )
    doc.add_heading("7. People, usability and responsibility", level=1)
    doc.add_heading("7.1 “Human error” is an incomplete diagnosis", level=2)
    doc.add_paragraph(
        "The original paper described employees as a major weakness and suggested that fear of mistakes could encourage organisations to replace people with automation. "
        "That framing now seems too blunt. A person may click a convincing message, reuse a password, postpone an update or work around a control; but those actions occur inside systems designed by organisations. "
        "If a single ordinary mistake can create catastrophic harm, the architecture, defaults, permissions, recovery process and organisational incentives also require scrutiny."
    )
    doc.add_paragraph(
        "Security guidance that imposes unlimited effort on users can be economically irrational from their perspective. "
        "Herley’s analysis of security advice argued that users bear the cost of precautions while many benefits are diffuse or invisible (2009). "
        "Usable security research likewise shows that controls must fit human capability and real tasks. The goal is not to eliminate human agency but to make the safe path comprehensible, convenient and forgiving."
    )
    add_labelled_callout(
        doc,
        "Design implication",
        "Treat mistakes as signals. Reduce privilege, provide phishing-resistant authentication, make reporting safe, isolate failures, log decisions, and design recovery so one action does not become an organisational disaster.",
    )
    doc.add_heading("7.2 Education still matters—but it is not a substitute for design", level=2)
    doc.add_paragraph(
        "The original project was right to value cyber education. Training helps people recognise suspicious behaviour, understand reporting routes and make informed choices. "
        "But awareness campaigns cannot repair insecure defaults, unsupported devices, unclear interfaces or absent leadership. "
        "Training must sit beside technical controls, clear ownership and a culture in which early reporting is rewarded rather than punished."
    )

    doc.add_heading("8. Automation and artificial intelligence", level=1)
    doc.add_heading("8.1 Automation is indispensable, not autonomous", level=2)
    doc.add_paragraph(
        "Automated asset discovery, configuration checking, telemetry analysis and response orchestration allow defenders to operate at a scale that manual work cannot match. "
        "However, automation reproduces its assumptions. Poor rules can block legitimate work, miss novel behaviour or create a false sense of control. "
        "Good systems therefore preserve auditability, allow proportionate human review and are tested against realistic failure modes."
    )
    doc.add_heading("8.2 AI changes both attack and defence", level=2)
    doc.add_paragraph(
        "The original paper presented AI primarily as a faster detector and asked whether it might replace cybersecurity staff. "
        "The 2026 position is more complex. AI can assist triage, anomaly detection, code analysis, summarisation and threat intelligence. "
        "It can also help attackers scale convincing content, search for weaknesses and automate parts of an intrusion. AI systems introduce their own risks through training data, model behaviour, integrations, access to sensitive information and over-trust in outputs."
    )
    doc.add_paragraph(
        "NIST’s AI Risk Management Framework organises responsible practice around governing, mapping, measuring and managing AI risk across the lifecycle (NIST, 2023; 2024). "
        "NCSC guidance similarly applies secure design, threat modelling, logging and incident management to machine-learning systems. "
        "The appropriate conclusion is not people versus machines. It is that accountable professionals should use automation while retaining judgement over context, consequence and exception."
    )

    doc.add_heading("9. What effective cybersecurity looks like", level=1)
    doc.add_paragraph(
        "A useful model combines the six NIST CSF 2.0 functions with secure-by-design practice. "
        "The functions are not a maturity score or a one-time sequence; together they provide a way to ask whether an organisation can make and sustain good risk decisions."
    )
    functions = [
        ("Govern", "Who owns the risk and sets priorities?", "Named responsibility, policy, risk appetite, supplier oversight"),
        ("Identify", "What do we depend on and what could go wrong?", "Asset inventory, dependency map, risk assessment, threat model"),
        ("Protect", "How is likely harm made harder?", "Secure defaults, MFA, least privilege, patching, backups, segmentation"),
        ("Detect", "How would we know something changed?", "Useful logs, monitoring, tested alerts, reporting routes"),
        ("Respond", "Can we contain and communicate?", "Rehearsed roles, isolation, decision logs, regulator and supplier routes"),
        ("Recover", "Can trusted service be restored?", "Tested backups, recovery objectives, learning review, improvement tracking"),
    ]
    for function, question, evidence in functions:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{function}. ")
        r.bold = True
        r.font.color.rgb = ACCENT_DARK
        p.add_run(f"{question} Evidence: {evidence}.")
    doc.add_paragraph(
        "A control should also be judged by whether people can use it correctly, whether failure is contained, and whether the organisation can demonstrate that it works.",
        style="Source note",
    )

    add_part_heading(
        doc,
        "IV",
        "Turn research into a product direction",
        "The work moves from a broad security question to a concrete tool for making device decisions clearer, safer and auditable.",
    )
    doc.add_heading("10. From research to a product direction", level=1)
    doc.add_paragraph(
        "The most durable question in the original EPQ was practical: how can security guidance reach the person making a real decision? "
        "My later undergraduate dissertation, “How do we decide what we provide?”, narrowed that broad concern to device procurement and provisioning. "
        "It examined how people choose connected devices, how security evidence can be compared, and how configuration advice can be turned into an operational process."
    )
    doc.add_paragraph(
        "That work led to the Device Provisioning Toolkit: a research prototype that discovers devices, compares relevant security signals and generates a tailored hardening guide. "
        "The product direction reflects the revised thesis of this paper. It does not promise perfect security. It tries to make a safer decision easier, document the reasoning, reduce insecure defaults and support repeatable action."
    )
    add_labelled_callout(
        doc,
        "Research-to-product thread",
        "EPQ: security is never finished → undergraduate research: device choices need clearer evidence → product prototype: turn security evidence into a usable, auditable provisioning workflow.",
    )

    doc.add_heading("11. Limitations", level=1)
    add_bullets(
        doc,
        (
            "The revised edition is a critical literature-based reappraisal, not a new empirical dissertation.",
            "The original Google Forms exercise cannot support population-level claims because of its size, recruitment method and inconsistent reported response counts.",
            "Official incident and survey statistics measure different populations and depend on detection and reporting; they should not be combined as if they were one dataset.",
            "The threat landscape changes quickly. Evidence is current to July 2026, but figures and guidance should be rechecked before later publication or recording.",
            "The Device Provisioning Toolkit is a research prototype. Retailer data, public vulnerability information and generated guidance can be incomplete and require review.",
        ),
    )

    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        "Cybersecurity cannot make modern digital life permanently free from danger, and describing it as either secure or insecure obscures the decisions that matter. "
        "The more defensible question is whether our institutions and products reduce likely harm, recognise failure, contain its consequences and recover in a trustworthy way."
    )
    doc.add_paragraph(
        "Current evidence shows both pressure and possibility. UK organisations continue to experience frequent attacks and nationally significant incidents have risen sharply. "
        "At the same time, authoritative frameworks describe practical controls, many attacks fail, and prepared organisations recover more effectively. "
        "The persistent weakness is implementation: too few organisations conduct risk assessments, assign board responsibility or maintain tested response plans."
    )
    doc.add_paragraph(
        "My answer in 2026 is therefore conditional. Contemporary cybersecurity can keep pace, but only when it is continuous, governed, usable and designed into products and services. "
        "Education and automation remain necessary, yet neither can carry responsibility alone. The best security makes the safer action easier, limits the cost of mistakes and turns evidence into decisions people can act on."
    )

    doc.add_heading("References", level=1)
    references = [
        (
            "CISA (2025).",
            "Product Security Bad Practices (updated guidance).",
            "https://www.cisa.gov/news-events/alerts/2025/01/17/cisa-and-fbi-release-updated-guidance-product-security-bad-practices",
        ),
        (
            "Department for Science, Innovation and Technology (DSIT) (2024).",
            "The UK Product Security and Telecommunications Infrastructure (Product Security) regime.",
            "https://www.gov.uk/government/publications/the-uk-product-security-and-telecommunications-infrastructure-product-security-regime",
        ),
        (
            "DSIT and Home Office (2026).",
            "Cyber Security Breaches Survey 2025/2026.",
            "https://www.gov.uk/government/statistics/cyber-security-breaches-survey-20252026/cyber-security-breaches-survey-20252026",
        ),
        (
            "DSIT and National Cyber Security Centre (NCSC) (2025).",
            "Cyber Governance Code of Practice.",
            "https://www.gov.uk/government/publications/cyber-governance-code-of-practice",
        ),
        (
            "European Union Agency for Cybersecurity (ENISA) (2025).",
            "ENISA Threat Landscape 2025.",
            "https://www.enisa.europa.eu/publications/enisa-threat-landscape-2025",
        ),
        (
            "Herley, C. (2009).",
            "So Long, and No Thanks for the Externalities: The Rational Rejection of Security Advice by Users.",
            "https://www.microsoft.com/en-us/research/publication/so-long-and-no-thanks-for-the-externalities-the-rational-rejection-of-security-advice-by-users/",
        ),
        (
            "NCSC (2018).",
            "Secure by Default.",
            "https://www.ncsc.gov.uk/information/secure-default",
        ),
        (
            "NCSC (2025a).",
            "Annual Review 2025 launch.",
            "https://www.ncsc.gov.uk/speech/annual-review-2025-richard-horne-speech",
        ),
        (
            "NCSC (2025b).",
            "NCSC Annual Review 2025: Incident management.",
            "https://www.ncsc.gov.uk/collection/ncsc-annual-review-2025/chapter-01-cyber-threat-to-the-uk/incident-management",
        ),
        (
            "NCSC (2025c).",
            "Software Security Code of Practice: Secure design and development.",
            "https://www.ncsc.gov.uk/collection/software-security-code-of-practice-implementation-guidance/theme-1-secure-design-development",
        ),
        (
            "NIST (2023).",
            "Artificial Intelligence Risk Management Framework (AI RMF 1.0).",
            "https://www.nist.gov/itl/ai-risk-management-framework",
        ),
        (
            "NIST (2024).",
            "Artificial Intelligence Risk Management Framework: Generative Artificial Intelligence Profile.",
            "https://nvlpubs.nist.gov/nistpubs/ai/NIST.AI.600-1.pdf",
        ),
        (
            "Pascoe, C., Quinn, S. and Scarfone, K. (2024).",
            "The NIST Cybersecurity Framework (CSF) 2.0.",
            "https://doi.org/10.6028/NIST.CSWP.29",
        ),
        (
            "Verizon (2025).",
            "2025 Data Breach Investigations Report.",
            "https://www.verizon.com/business/resources/reports/dbir/",
        ),
        (
            "von Solms, R. and van Niekerk, J. (2013).",
            "From information security to cyber security. Computers & Security, 38, 97–102.",
            "https://doi.org/10.1016/j.cose.2013.04.004",
        ),
    ]
    for ref in references:
        add_reference(doc, *ref)
    doc.add_paragraph(
        f"Reference review date: {date(2026, 7, 19).strftime('%-d %B %Y')}. "
        "Web links are included for transparency; access dates should be refreshed if the work is submitted or republished later.",
        style="Source note",
    )

    doc.add_heading("Appendix A. Revision at a glance", level=1)
    for theme, original, revised in (
        ("Question", "Binary security test", "Continuous risk and resilience test"),
        ("Evidence", "Mixed web sources", "Current official, standards and peer-reviewed sources"),
        ("People", "Primary weakness", "Participants in a designed system"),
        ("AI", "Possible replacement", "Governed support for accountable professionals"),
        ("Outcome", "Awareness as the answer", "Secure design, governance, usable controls and recovery"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{theme}: ")
        r.bold = True
        r.font.color.rgb = ACCENT_DARK
        p.add_run(f"{original} → {revised}.")

    path = OUT / "a-level-epq-dissertation.docx"
    doc.save(path)
    return path


def add_timing_heading(doc: Document, timing: str, title: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    r = p.add_run(f"{timing}  ")
    r.font.color.rgb = ACCENT_DARK
    p.add_run(title)


def add_direction(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph(style="Source note")
    r = p.add_run(f"{label.upper()}: ")
    r.bold = True
    r.font.color.rgb = ACCENT_DARK
    p.add_run(text)


def build_transcript() -> Path:
    doc = Document()
    title = "The question I asked before I knew how to ask it"
    configure_document(
        doc,
        title,
        "Research Notes video series, episode one production script",
    )
    add_title_page(
        doc,
        "Research Notes / Episode 01",
        title,
        "How my A-level cybersecurity EPQ became the starting point for products built around safer decisions",
        (
            "Presenter: Nathan Brown-Bennett",
            "Target running time: 10–12 minutes",
            "Production script · July 2026",
        ),
        "This script is designed for a direct-to-camera research video with document details, restrained data graphics and a short product bridge. "
        "Figures should be rechecked immediately before recording.",
    )

    doc.add_heading("Episode promise", level=1)
    doc.add_paragraph(
        "By the end, viewers should understand why “is cybersecurity secure?” is the wrong test, what the original A-level project got right and wrong, and how the corrected idea led towards a practical device-provisioning product."
    )
    add_labelled_callout(
        doc,
        "Core line",
        "I do not build security theatre. I build tools that make a safer decision easier to understand, repeat and defend.",
    )

    add_timing_heading(doc, "00:00–00:35", "Cold open")
    add_direction(doc, "Picture", "Direct to camera. Cut to the original EPQ title page, with the question highlighted.")
    doc.add_paragraph(
        "When I was doing my A-levels, I wrote my first dissertation around a question that sounded straightforward: "
        "is cybersecurity actually secure enough for modern-day use?"
    )
    doc.add_paragraph(
        "The question was imperfect. Cybersecurity is not a lock you can label secure or insecure. "
        "But hidden inside it was the question I have kept working on ever since: when technology becomes essential, how do we make protection practical enough for people to use—and strong enough to matter?"
    )
    doc.add_paragraph(
        "This is the first in a series about the research behind my products: what I believed, what the evidence now says, and how a dissertation became something you can actually use."
    )

    add_timing_heading(doc, "00:35–01:30", "Where the project began")
    add_direction(doc, "Picture", "Slow pan across the original contents page, survey charts and bibliography. Keep the 2020 date visible.")
    doc.add_paragraph(
        "The original project was written in 2020. I started with physical locks and the idea of perfect security, then moved through information security, cybersecurity, public awareness, human error, artificial intelligence and regulation."
    )
    doc.add_paragraph(
        "I also ran a small Google Forms quiz among people around me. It asked about hackers, public Wi-Fi, passwords and the differences between cyber and information security. "
        "It was useful to me as a student because the answers showed that familiar words did not produce a shared understanding."
    )
    doc.add_paragraph(
        "But it was not representative research. The sample was small, it was socially concentrated, and the write-up even contains inconsistent response counts. "
        "One question asked for a universal minimum password length, which was not a sound premise. "
        "So I keep it as part of the history of the work—not as a statistic about the public."
    )

    add_timing_heading(doc, "01:30–02:40", "What survived the rewrite")
    add_direction(doc, "Graphic", "Four words appear in sequence: No perfection. People matter. Automate carefully. Security is social.")
    doc.add_paragraph(
        "Four ideas survived. First, perfect security is not an operational goal. A control is not useless because a determined attacker may eventually get around it."
    )
    doc.add_paragraph(
        "Second, people matter. Security advice that nobody understands or can realistically follow is weak security."
    )
    doc.add_paragraph(
        "Third, automation is necessary. Modern environments create far more events, devices and dependencies than a person can manually check."
    )
    doc.add_paragraph(
        "And fourth, cybersecurity is social as well as technical. It touches education, business decisions, public services, product design and the way organisations respond when something goes wrong."
    )

    add_timing_heading(doc, "02:40–03:55", "What the security fields actually protect")
    add_direction(doc, "Graphic", "Use the four-field comparison from the deck, then highlight the same incident crossing all four boundaries.")
    doc.add_paragraph(
        "Before changing the question, I need to separate four terms that are often used as if they mean the same thing."
    )
    doc.add_paragraph(
        "Computer security protects an individual computing system: its operating system, software, memory, storage and local data. "
        "IT security protects an organisation’s technology estate—devices, identities, networks, cloud platforms, applications and backups."
    )
    doc.add_paragraph(
        "Information security protects information in any form, including paper and spoken information, not only digital files. "
        "Cybersecurity focuses on harm through connected digital systems, so it can include technology and information, but also people, services and physical processes reachable through them."
    )
    doc.add_paragraph(
        "The roles overlap because incidents cross boundaries. A stolen cloud credential can require endpoint evidence, IT access revocation, an information-impact assessment and a coordinated cyber response. "
        "The labels should clarify ownership and handoffs—not create silos."
    )

    add_timing_heading(doc, "03:55–05:10", "The question changes")
    add_direction(doc, "Graphic", "Cross out “Is cybersecurity secure?” Replace it with the revised research question.")
    doc.add_paragraph(
        "The biggest correction is the question itself. Security is not a permanent state. It is a decision about risk: what are we protecting, who might attack it, what could the harm be, and what can we reasonably do?"
    )
    doc.add_paragraph(
        "NIST’s current Cybersecurity Framework organises that work into six functions: govern, identify, protect, detect, respond and recover."
    )
    doc.add_paragraph(
        "That final pair changes the conversation. If your definition only asks whether an attack got through, every incident looks like total failure. "
        "If you also ask whether it was detected, contained and recovered from, you can distinguish exposure from catastrophe."
    )
    doc.add_paragraph(
        "So my revised question is this: can contemporary cybersecurity reduce and govern digital risk—and help us recover—at the pace at which technology and attackers change?"
    )

    add_timing_heading(doc, "05:10–06:50", "What the 2026 evidence says")
    add_direction(doc, "Graphic", "Use one number per frame. Always show the source and measurement period.")
    doc.add_paragraph(
        "The current evidence is serious. In the UK Cyber Security Breaches Survey for 2025 to 2026, 43 per cent of businesses said they had identified a breach or attack in the previous twelve months."
    )
    doc.add_paragraph(
        "The National Cyber Security Centre handled 204 nationally significant incidents in the year to August 2025—more than double the previous year’s 89."
    )
    doc.add_paragraph(
        "But here is the part that matters for the argument: only 30 per cent of businesses reported conducting a cyber-risk assessment, and only 25 per cent had a formal incident-response plan."
    )
    doc.add_paragraph(
        "Those figures do not prove that cybersecurity does not work. They show that proven practice is unevenly adopted. "
        "Some organisations are facing a modern threat environment with no tested plan for the moment prevention fails."
    )
    doc.add_paragraph(
        "And incident statistics have limits. An organisation with better monitoring may find more attacks. Some incidents are never detected or reported. "
        "The number matters—but it is not a simple league table for whether defence works."
    )

    add_timing_heading(doc, "06:50–08:10", "I would no longer blame the user")
    add_direction(doc, "Picture", "Show a hostile-looking security prompt, then simplify it into a clear, safe default.")
    doc.add_paragraph(
        "My original dissertation talked about human error as though people were mainly a vulnerability that automation might replace. I would not frame it that way now."
    )
    doc.add_paragraph(
        "People do make mistakes. But they make them inside systems that somebody designed. "
        "If one ordinary click can expose an entire organisation, we should ask about permissions, authentication, isolation, reporting culture and recovery—not only about the person who clicked."
    )
    doc.add_paragraph(
        "The better design question is: can we make the safe action easier? Can we use phishing-resistant authentication? Can we remove default passwords? "
        "Can somebody report a mistake early without worrying that honesty will punish them?"
    )
    doc.add_paragraph(
        "Training still matters. But training is not a patch for an unsafe product."
    )

    add_timing_heading(doc, "08:10–09:25", "AI does not remove accountability")
    add_direction(doc, "Graphic", "Split screen: AI-assisted detection on one side; AI-enabled phishing and model risk on the other.")
    doc.add_paragraph(
        "The original project also asked whether artificial intelligence could replace security staff. The better answer is that AI changes both sides of the problem."
    )
    doc.add_paragraph(
        "It can help defenders triage alerts, inspect code, summarise threat information and spot unusual patterns. "
        "It can also help attackers produce convincing messages, automate reconnaissance and work at greater scale."
    )
    doc.add_paragraph(
        "AI systems introduce their own dependencies and failure modes. Their data, permissions, integrations and outputs need to be governed and tested."
    )
    doc.add_paragraph(
        "So the useful model is not person versus machine. It is accountable people using automation, with evidence, audit trails and a way to challenge the result."
    )

    add_timing_heading(doc, "09:25–11:10", "How the research became a product")
    add_direction(doc, "Picture", "Show the undergraduate dissertation title, then current Device Provisioning Toolkit discovery and security-guide screens.")
    doc.add_paragraph(
        "The line from this first project to my later work is clearer in hindsight. The EPQ asked why protection is difficult to understand and apply."
    )
    doc.add_paragraph(
        "My undergraduate dissertation—“How do we decide what we provide?”—made that question specific. "
        "When an organisation buys or deploys a connected device, how does it compare the evidence? How does it avoid insecure defaults? How does advice become a repeatable workflow?"
    )
    doc.add_paragraph(
        "That work led to the Device Provisioning Toolkit. It is a research prototype that discovers devices, compares relevant security signals and produces a tailored hardening guide."
    )
    doc.add_paragraph(
        "It does not promise that a device is perfectly safe. That would repeat the mistake in my first question. "
        "It aims to make the decision more informed, the setup more consistent and the reasoning easier to audit."
    )
    doc.add_paragraph(
        "That is the product principle I have carried forward: do not merely tell people to be more secure. Build what helps them take the next safer action."
    )

    add_timing_heading(doc, "11:10–12:00", "Answer and close")
    add_direction(doc, "Picture", "Return to camera. End on a simple research lineage graphic and the next episode title.")
    doc.add_paragraph(
        "Can cybersecurity keep pace with modern digital risk? My answer is conditional."
    )
    doc.add_paragraph(
        "Yes—when it is continuous, governed, designed into the product, usable by real people and prepared for recovery."
    )
    doc.add_paragraph(
        "No—when it is treated as antivirus, an annual training slide, a checklist nobody tests, or a problem delegated to one technical team."
    )
    doc.add_paragraph(
        "The question I asked at A-level was not quite the right one. But learning why it was wrong gave me a much better direction."
    )
    doc.add_paragraph(
        "In the next episode, I’ll move into my undergraduate dissertation and show how that broad argument became a working comparison and device-provisioning system."
    )

    doc.add_page_break()
    doc.add_heading("Production notes", level=1)
    add_bullets(
        doc,
        (
            "Tone: reflective and precise, not embarrassed by the early work. The point is visible intellectual development.",
            "Data graphics: display the denominator, period and source; do not animate unrelated numbers together.",
            "Original document: label every excerpt “2020 A-level submission”.",
            "Revised document: label every excerpt “2026 critical edition”.",
            "Product footage: use real current screens. Avoid implying that the prototype certifies a device or guarantees procurement outcomes.",
            "Accessibility: add captions, describe charts in speech, avoid fast text transitions, and provide the transcript on the website.",
        ),
    )
    doc.add_heading("On-screen source list", level=1)
    for ref in (
        ("DSIT & Home Office (2026).", "Cyber Security Breaches Survey 2025/2026.", "https://www.gov.uk/government/statistics/cyber-security-breaches-survey-20252026"),
        ("NCSC (2025).", "Annual Review 2025.", "https://www.ncsc.gov.uk/section/annual-review-2025"),
        ("NIST (2024).", "Cybersecurity Framework 2.0.", "https://doi.org/10.6028/NIST.CSWP.29"),
        ("ENISA (2025).", "Threat Landscape 2025.", "https://www.enisa.europa.eu/publications/enisa-threat-landscape-2025"),
        ("NCSC (2025).", "Software Security Code of Practice.", "https://www.ncsc.gov.uk/collection/software-security-code-of-practice-implementation-guidance"),
    ):
        add_reference(doc, *ref)

    path = OUT / "a-level-epq-video-transcript.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    dissertation = build_dissertation()
    transcript = build_transcript()
    print(dissertation)
    print(transcript)
