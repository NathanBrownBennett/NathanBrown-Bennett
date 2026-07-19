#!/usr/bin/env python3
"""Build the full 2026 standalone EPQ dissertation.

The document uses a narrative-proposal style system with an editorial cover.
The front matter is single-column; the dissertation body, references and
appendices use a professional two-column academic layout.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from build_epq_research_package import (
    ACCENT_DARK,
    DISPLAY,
    FONT,
    INK,
    MUTED,
    add_bullets,
    add_hyperlink,
    add_labelled_callout,
    add_numbered,
    configure_document,
    set_section_columns,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "assets" / "documents"
OUTPUT = OUT / "a-level-epq-dissertation.docx"


def add_para(doc: Document, text: str, *, lead: bool = False) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.first_line_indent = Cm(0 if lead else 0.42)
    if lead and paragraph.runs:
        paragraph.runs[0].font.size = Pt(10.4)


def add_section_intro(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_after = Pt(7)
    if paragraph.runs:
        paragraph.runs[0].font.color.rgb = MUTED
        paragraph.runs[0].italic = True


def add_reference(doc: Document, citation: str, title: str, url: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.55)
    paragraph.paragraph_format.first_line_indent = Cm(-0.55)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.add_run(citation + " ")
    add_hyperlink(paragraph, title, url)


def add_restarted_numbered(doc: Document, items: tuple[str, ...]) -> None:
    """Add a true Word numbered list whose numbering restarts at one."""
    numbering = doc.part.numbering_part.element
    base_num_id = int(doc.styles["List Number"].element.pPr.numPr.numId.val)
    base_num = numbering.xpath(f'./w:num[@w:numId="{base_num_id}"]')[0]
    abstract_num_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    new_num_id = max(existing_ids, default=0) + 1

    new_num = OxmlElement("w:num")
    new_num.set(qn("w:numId"), str(new_num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_num_id)
    new_num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    new_num.append(level_override)
    numbering.append(new_num)

    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
        num_pr.get_or_add_ilvl().val = 0
        num_pr.get_or_add_numId().val = new_num_id
        paragraph.add_run(item)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph("")
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("EXTENDED PROJECT QUALIFICATION  /  UPDATED DISSERTATION")
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(9)
    run.font.color.rgb = ACCENT_DARK

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(14)
    run = title.add_run("Can contemporary cybersecurity provide effective security for modern digital life?")
    run.bold = True
    run.font.name = DISPLAY
    run.font.size = Pt(29)
    run.font.color.rgb = INK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(42)
    run = subtitle.add_run(
        "An investigation of what security means, how computer, IT, information and "
        "cybersecurity divide responsibility, and whether their combined practice can keep pace with modern risk"
    )
    run.font.name = FONT
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED

    for line, bold in (
        ("Nathan Brown-Bennett", True),
        ("Independent updated edition based on the 2020 A-level EPQ", False),
        ("Research current to July 2026", False),
    ):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(line)
        run.bold = bold
        run.font.name = FONT
        run.font.size = Pt(10.5 if bold else 9.5)
        run.font.color.rgb = INK if bold else MUTED

    doc.add_paragraph("")
    statement = doc.add_paragraph()
    statement.alignment = WD_ALIGN_PARAGRAPH.CENTER
    statement.paragraph_format.left_indent = Cm(2.4)
    statement.paragraph_format.right_indent = Cm(2.4)
    run = statement.add_run(
        "This is a complete standalone dissertation. The original assessed submission is preserved "
        "separately and is used only as the historical starting point for the updated research question."
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    doc.add_page_break()


def add_front_matter(doc: Document) -> None:
    doc.add_heading("Abstract", level=1)
    add_para(
        doc,
        "Modern societies rely on computers, networks, cloud platforms, connected products and digital "
        "information for activities ranging from personal communication to healthcare and national "
        "infrastructure. This dependence creates a difficult question: if cyber incidents remain common, "
        "can cybersecurity reasonably be described as providing effective security? This dissertation "
        "argues that the question cannot be answered by counting attacks or asking whether a system is "
        "simply secure or insecure. Security is a managed condition in which unacceptable harm is reduced "
        "through proportionate controls, clear responsibility, evidence and recovery. It is neither a "
        "permanent technical state nor a guarantee that failure will never occur.",
        lead=True,
    )
    add_para(
        doc,
        "The investigation first separates four overlapping fields. Computer security concentrates on an "
        "individual computing system and the data it processes. IT security protects an organisation's "
        "technology estate and the services needed to operate it. Information security protects information "
        "in any form throughout its lifecycle. Cybersecurity addresses harm created through interconnected "
        "digital systems, including effects on people, organisations and physical services. These boundaries "
        "are analytical rather than absolute: real incidents move between them, which makes ownership and "
        "handoffs as important as specialist expertise.",
    )
    add_para(
        doc,
        "A structured review of current government, standards, professional and peer-reviewed sources is "
        "used to examine threat prevalence, governance, human factors, secure design, automation, artificial "
        "intelligence, response and resilience. UK evidence shows both continuing exposure and uneven "
        "preparedness. In the Cyber Security Breaches Survey 2025/2026, 43 per cent of businesses identified "
        "a breach or attack in the preceding year, while only 30 per cent reported a cyber-risk assessment "
        "and 25 per cent had a formal incident-response plan. The National Cyber Security Centre handled "
        "204 nationally significant incidents in 2024-25. These figures show serious risk, but they do not "
        "prove that cybersecurity is futile: detection increases reported prevalence, many attempted attacks "
        "are prevented, and mature response limits consequence.",
    )
    add_para(
        doc,
        "The dissertation concludes that contemporary cybersecurity can provide effective security, but only "
        "conditionally. It works when security is governed as organisational risk, built into products and "
        "services, usable by real people, monitored continuously and prepared for recovery. It fails when "
        "organisations purchase isolated controls, transfer responsibility to users, leave ownership unclear "
        "or measure success only by the absence of visible incidents. The most defensible goal is therefore "
        "not perfect protection, but a demonstrable capacity to prevent avoidable harm, detect change, "
        "contain failure and restore trusted operation.",
    )

    doc.add_heading("Research question", level=1)
    add_labelled_callout(
        doc,
        "Primary question",
        "Can contemporary cybersecurity provide effective security for modern digital life, and how do the "
        "responsibilities of computer security, IT security, information security and cybersecurity combine "
        "to produce that protection?",
    )

    doc.add_heading("Contents", level=1)
    contents = (
        "Introduction",
        "Methodology and evaluation approach",
        "What security means",
        "The four security fields: scope, responsibility and boundaries",
        "Security work: roles, outputs and handoffs",
        "The modern threat environment",
        "Does cybersecurity work?",
        "People, usability and security culture",
        "Technical architecture and secure-by-design practice",
        "Automation and artificial intelligence",
        "Governance, law and ethical limits",
        "Integrated incident analysis",
        "Discussion",
        "Conclusion",
    )
    add_numbered(doc, contents)
    for item in ("References", "Appendices"):
        paragraph = doc.add_paragraph(item)
        paragraph.paragraph_format.left_indent = Cm(0.65)
        paragraph.paragraph_format.space_after = Pt(3)

    doc.add_heading("Research integrity statement", level=1)
    add_para(
        doc,
        "This dissertation is an independently produced 2026 investigation developed from the question first "
        "considered in my A-level EPQ. It is not presented as the assessed 2020 submission and does not alter that "
        "historical record. The new work has a different research question, method, structure, evidence base and "
        "conclusion. The original questionnaire is excluded from the evidential argument because its sampling and "
        "internal inconsistencies do not support general claims. Where the earlier project suggested a useful theme, "
        "that theme has been researched again rather than copied forward as fact."
    )
    add_para(
        doc,
        "Claims in this edition are attributed through in-text author-date references and a linked reference list. "
        "Statistics are stated with their reporting period and population wherever they materially affect the "
        "argument. Official publications are not treated as infallible: their methods, reporting thresholds and "
        "limitations are considered in Chapter 2. Industry evidence is used selectively and is not allowed to "
        "override more transparent public data merely because it reports a larger number."
    )

    doc.add_heading("Terminology convention", level=1)
    add_para(
        doc,
        "Security terminology is contested and context-dependent. This dissertation therefore uses four working "
        "definitions rather than claiming that every standard or employer draws the same boundary. Computer "
        "security refers to protection of an individual computing system. IT security refers to protection of an "
        "organisation's technology estate and operational services. Information security refers to protection of "
        "information in any form and the obligations attached to it. Cybersecurity refers to the management of "
        "harm arising through interconnected digital systems, including consequences for people, organisations "
        "and physical services. Chapter 4 explains and tests these definitions."
    )
    add_para(
        doc,
        "The word breach follows the wording of the source being discussed and does not always mean confirmed "
        "loss of data. Attack refers to intentional activity intended to defeat or misuse a system; incident refers "
        "to an event requiring assessment or response. Risk is used as a contextual judgement about uncertainty "
        "and consequence, not as a synonym for threat or vulnerability. Effective security means a defensible "
        "reduction in unacceptable harm, supported by evidence and recovery, rather than the total absence of "
        "malicious activity."
    )

    doc.add_heading("Summary of the argument", level=1)
    add_para(
        doc,
        "The dissertation develops one cumulative argument. First, absolute security is not an operationally useful "
        "standard for complex and changing systems. Second, the four security fields provide complementary lenses "
        "whose value lies in clear ownership and handoffs. Third, contemporary evidence shows substantial exposure "
        "but also a persistent gap between available practice and consistent adoption. Fourth, people, technology, "
        "governance and suppliers form one security system; transferring responsibility to any single part produces "
        "fragile protection. Finally, cybersecurity can be effective when it is continuous, proportionate, usable "
        "and designed for recovery. Each chapter supplies a necessary part of that conclusion."
    )
    doc.add_page_break()


def add_chapter_1(doc: Document) -> None:
    doc.add_heading("1. Introduction", level=1)
    add_section_intro(
        doc,
        "This chapter establishes the problem, research question, scope and criteria by which the answer will be judged."
    )
    doc.add_heading("1.1 Background", level=2)
    add_para(
        doc,
        "Security has always concerned the protection of valued things from unacceptable harm. A lock protects "
        "access to a room; a legal process protects rights and evidence; a safety system limits physical injury. "
        "Digital technology changes the objects, scale and speed of protection, but not the underlying need to "
        "decide what matters, what could damage it and what response is proportionate. Computers now mediate "
        "banking, employment, education, entertainment, transport, healthcare, public administration and "
        "personal relationships. The same systems that create convenience also concentrate dependence. A stolen "
        "credential can expose years of information; a software defect can be reproduced across thousands of "
        "organisations; disruption to a connected service can create consequences outside the computer itself."
    )
    add_para(
        doc,
        "Public discussion often responds to this tension with two incompatible claims. The first is that "
        "cybersecurity is increasingly important because threats are growing. The second is that cybersecurity "
        "must be failing because attacks still occur. Both contain part of the truth, but neither supplies a "
        "sound test of effectiveness. Medicine is not ineffective because illness exists, and fire safety is "
        "not worthless because buildings can burn. A protective discipline must instead be judged against its "
        "purpose: whether it reduces foreseeable harm, makes failure less likely, discovers problems early, "
        "limits their consequence and supports recovery."
    )
    add_para(
        doc,
        "The terminology also creates confusion. Computer security, IT security, information security and "
        "cybersecurity are frequently used as synonyms. Standards themselves contain overlapping definitions, "
        "and job titles vary between employers. However, the terms draw attention to different protected "
        "objects and different decisions. A security engineer hardening an operating system, an identity "
        "administrator controlling staff access, an information-security manager assessing regulatory impact "
        "and an incident responder containing malicious activity may all contribute to the same event without "
        "performing the same work. Treating every task as generic 'cyber' work hides ownership and encourages "
        "gaps between teams."
    )

    doc.add_heading("1.2 Aim and objectives", level=2)
    add_para(
        doc,
        "The aim is to determine whether contemporary cybersecurity can provide effective protection for "
        "modern digital life without relying on an unrealistic promise of perfect security. The investigation "
        "has six objectives:"
    )
    add_bullets(
        doc,
        (
            "define security as a risk, control and resilience problem rather than a binary label;",
            "distinguish computer, IT, information and cybersecurity by what each protects and the decisions each owns;",
            "identify representative roles, work products and handoffs across those fields;",
            "evaluate current evidence about threats, breaches, preparedness and recovery;",
            "examine how human factors, technical design, automation, governance and law affect outcomes; and",
            "construct a reasoned test for deciding when cybersecurity is effective.",
        ),
    )

    doc.add_heading("1.3 Scope", level=2)
    add_para(
        doc,
        "The dissertation concentrates on civilian digital security in the United Kingdom while using "
        "international standards and research where they provide a stronger conceptual basis. It considers "
        "individual computers, organisational IT, information handling, internet-connected products, cloud "
        "services and the human and governance systems surrounding them. National military cyber operations, "
        "offensive cyber capability and classified intelligence are outside scope except where public sources "
        "help explain the broader threat environment."
    )
    add_para(
        doc,
        "The phrase 'modern digital life' includes both personal and organisational dependence. The evidence "
        "base is stronger for organisations because governments and industry publish survey and incident data "
        "at that level. Consumer experience is therefore discussed through product security, identity, privacy "
        "and connected-device examples rather than treated as statistically equivalent to business data. The "
        "research is current to July 2026; fast-moving figures and guidance will require later review."
    )

    doc.add_heading("1.4 Standard for an answer", level=2)
    add_para(
        doc,
        "An answer of 'yes' cannot mean that every attack is prevented. It must mean that suitable combinations "
        "of governance, people, process and technology can measurably reduce risk and consequence when they are "
        "properly selected and maintained. An answer of 'no' would require evidence that available practices "
        "cannot materially improve those outcomes, or that their cost and complexity make them unusable in the "
        "environments they are intended to protect. The conclusion will therefore distinguish the capability "
        "of the discipline from the quality of its implementation."
    )


def add_chapter_2(doc: Document) -> None:
    doc.add_heading("2. Methodology and evaluation approach", level=1)
    add_section_intro(
        doc,
        "The investigation is a structured, critical literature review supported by official statistics and an applied incident analysis."
    )
    doc.add_heading("2.1 Research design", level=2)
    add_para(
        doc,
        "A literature-based method is appropriate because the question concerns the boundaries and effectiveness "
        "of several disciplines rather than the behaviour of one product or population. The research combines "
        "four types of evidence. Standards and glossaries establish terminology and recognised practices. "
        "Government surveys and incident reports describe exposure and preparedness. Peer-reviewed and "
        "professional research examines conceptual boundaries and human behaviour. Regulatory and technical "
        "guidance shows how responsibility is changing in practice."
    )
    add_para(
        doc,
        "Sources were prioritised in the following order: current primary publications from NIST, NCSC, UK "
        "government and ENISA; legislation and official implementation guidance; peer-reviewed research with a "
        "clear method or durable conceptual contribution; and major industry reports where their datasets add "
        "detail not available from public statistics. News articles, vendor marketing pages and unsourced "
        "explainers were not used as evidence for central claims. Older research was retained only where it "
        "introduced an idea that remains relevant, such as the rational cost of security advice to users."
    )

    doc.add_heading("2.2 Search and selection", level=2)
    add_para(
        doc,
        "Searches covered definitions of computer, information, IT and cybersecurity; security risk and "
        "resilience; cyber workforce roles; UK breach prevalence and incident response; secure-by-design "
        "engineering; usable security; artificial-intelligence risk; product security regulation; and "
        "governance. Sources were checked for publication date, responsible organisation, methodology, scope "
        "and whether the claim being used could be traced to the original publication. Where a current webpage "
        "summarised a formal standard, the standard itself was preferred."
    )
    add_para(
        doc,
        "The NIST glossary illustrates why context matters. It aggregates definitions from different documents "
        "rather than declaring a single preferred definition for every term. Some sources use computer security, "
        "IT security and cybersecurity as synonyms; other research treats cybersecurity as a wider social and "
        "organisational domain. The method therefore does not pretend that one universal taxonomy exists. It "
        "constructs working definitions, states their purpose and tests whether they clarify responsibility."
    )

    doc.add_heading("2.3 Interpreting incident statistics", level=2)
    add_para(
        doc,
        "Breach statistics are not a direct score for defensive effectiveness. An organisation with mature "
        "monitoring may identify more attacks than one that cannot see them. Surveys depend on respondent "
        "knowledge and willingness to disclose; administrative incident data records only events that reach the "
        "reporting body; categories and thresholds change over time. Larger organisations have more assets and "
        "may be targeted more often, but they also tend to have stronger detection. Comparisons must therefore "
        "consider denominator, period, population, severity and measurement method."
    )
    add_para(
        doc,
        "The Cyber Security Breaches Survey 2025/2026 is valuable because it publishes its sampling approach, "
        "question bases and statistical-reliability guidance. Its estimate that 43 per cent of businesses "
        "identified a breach or attack describes recognised exposure during a twelve-month period; it does not "
        "claim that 57 per cent experienced none. Similarly, the NCSC's 204 nationally significant incidents "
        "describe cases meeting its threshold and handled during the reporting year, not every UK incident."
    )

    doc.add_heading("2.4 Limitations and ethics", level=2)
    add_para(
        doc,
        "No new participant data was collected. The small questionnaire in the 2020 submission is not reused as "
        "evidence because its convenience sample, inconsistent response counts and leading questions cannot "
        "support claims about the public. This removes an empirical component but improves the integrity of the "
        "updated dissertation. Future work could test the proposed field definitions and responsibility model "
        "through interviews with practitioners from security operations, IT, risk, privacy and product teams."
    )
    add_para(
        doc,
        "The research discusses defensive concepts and publicly documented threats without providing operational "
        "instructions for intrusion. Ethical evaluation is part of the subject rather than an afterthought: "
        "security controls can monitor people, restrict access and process sensitive data, so effectiveness must "
        "include legality, necessity, proportionality and accountability. A control that reduces one technical "
        "risk by creating unjustified surveillance is not automatically a good security outcome."
    )


def add_chapter_3(doc: Document) -> None:
    doc.add_heading("3. What security means", level=1)
    add_section_intro(
        doc,
        "Security begins with valued assets and unacceptable harm. Controls are useful only in relation to that context."
    )
    doc.add_heading("3.1 From absolute safety to managed risk", level=2)
    add_para(
        doc,
        "In everyday language, secure can imply freedom from danger. That absolute meaning becomes unstable when "
        "applied to complex systems. A computer contains hardware, firmware, operating-system components, "
        "applications, accounts, networks and external services, each changing over time. Its behaviour depends "
        "on administrators, developers, users and suppliers. No finite assessment can prove that every component "
        "is free from unknown weakness or that every future attacker will fail."
    )
    add_para(
        doc,
        "Operational security is better understood as the managed reduction of unacceptable harm. The process "
        "identifies assets, threats, vulnerabilities, likelihood and consequence; selects proportionate controls; "
        "and reviews whether residual risk is acceptable. This does not weaken the meaning of security. It makes "
        "the claim testable. A hospital and a home media player should not apply identical controls, but both "
        "should be able to explain what they protect, which failure matters and why their chosen protection is "
        "reasonable."
    )
    add_labelled_callout(
        doc,
        "Working definition",
        "Security is the justified and continuously reviewed reduction of unacceptable harm to valued assets, "
        "supported by controls that prevent, detect, contain and recover from failure.",
    )

    doc.add_heading("3.2 Assets, threats, vulnerabilities and controls", level=2)
    add_para(
        doc,
        "An asset is anything whose loss, exposure, alteration or unavailability would matter. Assets include "
        "devices, data, identities, money, intellectual property, services, reputation, safety and public trust. "
        "A threat is a circumstance or actor capable of causing harm. A vulnerability is a weakness that makes "
        "that harm possible. Risk combines uncertainty about occurrence with consequence. A control changes the "
        "likelihood, impact or detectability of an unwanted event."
    )
    add_para(
        doc,
        "This vocabulary prevents category errors. Malware is not itself a risk statement; it is one possible "
        "threat mechanism. 'Install antivirus' is not a complete strategy; it is a control whose value depends on "
        "coverage, configuration, telemetry, response and the threats under consideration. Likewise, a "
        "vulnerability score cannot decide business priority without knowing where the affected component is "
        "used, whether it is exposed, what data or service depends on it and what compensating controls exist."
    )

    doc.add_heading("3.3 Security properties", level=2)
    add_para(
        doc,
        "Information-security practice commonly begins with confidentiality, integrity and availability. "
        "Confidentiality limits disclosure to authorised people and processes. Integrity protects accuracy, "
        "completeness and authorised change. Availability ensures that authorised users can access information "
        "and services when required. The three properties interact: encrypting information may protect "
        "confidentiality, but inaccessible keys can destroy availability; unrestricted availability may expose "
        "confidentiality; an available service returning manipulated results lacks integrity."
    )
    add_para(
        doc,
        "Modern systems require additional properties. Authenticity supports confidence that an identity, message "
        "or component is genuine. Accountability links actions to responsible entities through evidence such as "
        "logs and approvals. Non-repudiation makes it difficult for a party to deny a relevant action. Privacy "
        "concerns appropriate processing of information about people, not merely secrecy. Safety addresses "
        "physical injury and environmental harm. Resilience concerns the ability to anticipate, withstand, "
        "recover and adapt when disruption occurs. These properties explain why a control can succeed technically "
        "while the overall system still fails its users."
    )

    doc.add_heading("3.4 Prevention, detection, response and recovery", level=2)
    add_para(
        doc,
        "Preventive controls reduce the chance that harmful action succeeds: secure configuration, authentication, "
        "segmentation, input validation and patching are examples. Detective controls reveal relevant change "
        "through logs, monitoring, alerting and reporting. Responsive controls contain harm, preserve evidence, "
        "communicate decisions and remove an attacker. Recovery restores trusted operation and incorporates "
        "learning. Overinvestment in prevention can create brittle security if detection and recovery are absent."
    )
    add_para(
        doc,
        "NIST CSF 2.0 expresses a similar lifecycle through Govern, Identify, Protect, Detect, Respond and Recover. "
        "The addition and prominence of Govern matters because technology cannot decide risk appetite, allocate "
        "accountability or balance competing obligations. The functions are concurrent rather than a one-off "
        "checklist: an organisation continuously learns about dependencies, changes protection and improves its "
        "ability to respond."
    )


def add_chapter_4(doc: Document) -> None:
    doc.add_heading("4. The four security fields", level=1)
    add_section_intro(
        doc,
        "The fields overlap, but each begins from a different protected object and operational boundary."
    )
    doc.add_heading("4.1 Computer security", level=2)
    add_para(
        doc,
        "Computer security protects an individual computing system and the information processed or stored by it. "
        "The system may be a laptop, server, mobile phone, embedded controller or virtual machine. Its boundary "
        "includes hardware trust, firmware, boot process, operating system, applications, memory, local storage, "
        "device interfaces and local accounts. Typical questions ask whether code can execute without authority, "
        "whether one process can interfere with another, whether stored data is protected and whether the system "
        "can be restored to a trusted state."
    )
    add_para(
        doc,
        "Relevant controls include secure boot, code signing, operating-system hardening, memory protection, "
        "application sandboxing, endpoint detection, disk encryption, local firewalling, vulnerability management "
        "and controlled administrative privilege. Computer security is narrower than organisational IT security "
        "because a device can be well hardened while identities, networks, backups or cloud configuration remain "
        "weak. It is also foundational: wider security depends on endpoints behaving predictably."
    )

    doc.add_heading("4.2 IT security", level=2)
    add_para(
        doc,
        "IT security protects the technology estate used to deliver an organisation's operations. It includes "
        "endpoints but extends to directories, identities, networks, email, collaboration platforms, business "
        "applications, databases, cloud services, backups, service-management processes and supplier connections. "
        "The central concern is dependable and controlled operation of technology at organisational scale."
    )
    add_para(
        doc,
        "IT-security decisions include how users join and leave, which roles receive access, how systems are "
        "configured, how changes are approved, how vulnerabilities are prioritised, what is logged, how backups "
        "are protected and how service is restored. The field is closely connected to IT operations. Separation "
        "of duties is useful, but an antagonistic boundary between 'IT' and 'security' is counterproductive "
        "because secure outcomes depend on routine administration being performed consistently."
    )

    doc.add_heading("4.3 Information security", level=2)
    add_para(
        doc,
        "Information security protects information and its value throughout its lifecycle, regardless of format. "
        "A printed contract, spoken medical detail, paper visitor list and encrypted database can all contain "
        "information requiring protection. The field asks how information is created, classified, accessed, "
        "shared, retained, transformed and destroyed, and which legal, contractual or organisational obligations "
        "apply."
    )
    add_para(
        doc,
        "Its scope therefore includes policies, risk assessment, classification, records management, supplier "
        "assurance, physical handling, business continuity and audit as well as technical controls. Information "
        "security provides a durable link between technology and business meaning. A database administrator may "
        "know that a table is encrypted; the information owner must determine whether the data should have been "
        "collected, who needs it, how long it should remain and what disclosure would mean."
    )

    doc.add_heading("4.4 Cybersecurity", level=2)
    add_para(
        doc,
        "Cybersecurity addresses harm arising through interconnected digital systems. It includes the protection "
        "of technology and information but gives greater attention to adversarial activity, networks, external "
        "dependencies, digital services and consequences extending into society or the physical world. Von Solms "
        "and van Niekerk argue that cybersecurity extends beyond information security because not all assets in "
        "cyberspace are information: people, infrastructure and societal functions may also be harmed."
    )
    add_para(
        doc,
        "This broader scope includes threat intelligence, security operations, incident coordination, online "
        "fraud, connected-product security, operational technology and national resilience. Cybersecurity does "
        "not replace the other fields. It connects them when the source or path of harm is digital and networked. "
        "For example, ransomware involves endpoint behaviour, IT-wide identity and backup decisions, information "
        "impact, business continuity, law enforcement and sometimes physical service disruption."
    )

    doc.add_heading("4.5 Adjacent fields", level=2)
    add_para(
        doc,
        "Privacy, data protection, physical security, safety, fraud, resilience and digital forensics intersect "
        "with the four fields. Privacy asks whether personal information is processed fairly and appropriately; "
        "security supplies some protections but cannot establish lawful purpose by itself. Safety concentrates on "
        "preventing injury, which can conflict with confidentiality or access restrictions during emergencies. "
        "Fraud teams focus on deceptive transactions and financial loss. Digital forensics preserves and analyses "
        "evidence. These domains should coordinate without being collapsed into one undefined security function."
    )

    doc.add_heading("4.6 Where the boundaries overlap", level=2)
    add_para(
        doc,
        "The four-field model is best understood as a set of lenses. Computer security asks whether a device can "
        "be trusted. IT security asks whether the technology estate is controlled and dependable. Information "
        "security asks whether information and obligations are protected. Cybersecurity asks how connected "
        "digital activity can cause wider harm and how that activity should be anticipated and answered. A mature "
        "organisation applies all four questions to important systems."
    )
    add_labelled_callout(
        doc,
        "Boundary principle",
        "Labels should clarify ownership, expertise and handoffs. They become harmful when they create silos or "
        "allow one team to assume that another has accepted a risk.",
    )


def add_chapter_5(doc: Document) -> None:
    doc.add_heading("5. Security work: roles, outputs and handoffs", level=1)
    add_section_intro(
        doc,
        "Job titles are inconsistent; work is better understood through responsibilities, decisions and evidence."
    )
    doc.add_heading("5.1 Why roles are difficult to classify", level=2)
    add_para(
        doc,
        "A 'security engineer' may design cloud controls in one organisation, manage endpoint tools in another "
        "and test applications in a third. A systems administrator may perform security-critical work without "
        "holding a security title. The NICE Workforce Framework addresses this problem by describing work roles "
        "through tasks, knowledge and skills rather than assuming that job titles are universal. Its current "
        "components provide a common language for workforce planning while recognising that one job can combine "
        "several roles and that work roles are not identical to positions."
    )

    doc.add_heading("5.2 Roles centred on computer security", level=2)
    add_para(
        doc,
        "Endpoint-security engineers, platform-security engineers, operating-system specialists, malware analysts "
        "and vulnerability researchers concentrate on the behaviour of individual systems. Their outputs include "
        "hardened baselines, secure-build configurations, endpoint policies, vulnerability findings, exploitability "
        "assessments, malware reports and evidence that a device can be rebuilt or recovered. Software-security "
        "engineers and application-security testers overlap when the protected computer hosts code they evaluate."
    )
    add_para(
        doc,
        "These roles need deep technical access but should not decide business consequence alone. A critical "
        "vulnerability on an isolated laboratory device may be less urgent than a moderate weakness on an "
        "internet-facing identity server. Computer-security evidence must be handed to IT owners and risk owners "
        "who understand exposure, dependency and operational impact."
    )

    doc.add_heading("5.3 Roles centred on IT security", level=2)
    add_para(
        doc,
        "Security administrators, identity and access specialists, network-security engineers, cloud-security "
        "engineers, vulnerability managers, backup specialists and security operations analysts protect the "
        "organisational estate. They operate access reviews, segmentation, configuration standards, patch "
        "programmes, monitoring platforms, email controls, secrets management and recovery systems. Their outputs "
        "include inventories, access records, change evidence, alert rules, remediation plans and service-recovery "
        "tests."
    )
    add_para(
        doc,
        "IT-security roles work closest to routine operations. That position gives them practical knowledge of "
        "technical debt, maintenance windows, supplier constraints and user needs, but can also create conflicts "
        "between availability and change. Governance must establish how urgent risk decisions are escalated so "
        "that administrators are not forced to accept business risk informally."
    )

    doc.add_heading("5.4 Roles centred on information security", level=2)
    add_para(
        doc,
        "Information-security managers, risk analysts, governance and compliance specialists, security architects, "
        "auditors and information owners translate organisational obligations into requirements. Their work "
        "includes risk registers, policies, control frameworks, classification schemes, supplier assessments, "
        "assurance reviews, continuity requirements and reports to leadership. Data-protection professionals "
        "overlap but retain a distinct focus on lawful and fair processing of personal information."
    )
    add_para(
        doc,
        "These roles can fail if governance becomes documentation detached from operations. A policy stating that "
        "access is reviewed annually has little value if the identity system cannot produce reliable membership "
        "data. Effective information security therefore asks for evidence from technical teams and tests whether "
        "the control changes real exposure, not merely whether a document exists."
    )

    doc.add_heading("5.5 Roles centred on cybersecurity", level=2)
    add_para(
        doc,
        "SOC analysts, threat hunters, incident responders, threat-intelligence analysts, penetration testers, "
        "digital-forensics specialists, cybercrime investigators and operational-technology security engineers "
        "focus on adversarial activity and connected harm. Their outputs include alerts, incident timelines, "
        "containment actions, intelligence assessments, attack-path findings, forensic images, lessons learned "
        "and coordinated communications."
    )
    add_para(
        doc,
        "Incident response is inherently cross-functional. Responders may need authority to isolate devices, reset "
        "identities, interrupt services, preserve evidence, notify regulators, contact suppliers and communicate "
        "with affected people. Technical competence without pre-agreed decision rights creates delay at the moment "
        "when delay is most costly."
    )

    doc.add_heading("5.6 Leadership and shared responsibility", level=2)
    add_para(
        doc,
        "Boards and senior leaders own organisational risk even when specialists manage controls. Product managers "
        "and developers own design decisions; procurement teams influence supplier exposure; HR manages joining, "
        "movement and leaving; legal teams interpret obligations; communications teams manage public information; "
        "and every user operates within the system. 'Security is everyone's responsibility' is useful only if it "
        "does not erase accountable owners. Shared participation must sit beneath named decision authority."
    )
    add_para(
        doc,
        "A good handoff records the asset, event, evidence, uncertainty, owner, decision and deadline. For example, "
        "a vulnerability analyst establishes technical exposure; the service owner explains operational context; "
        "the information owner describes affected data; the risk owner accepts, transfers, mitigates or avoids the "
        "risk; and assurance checks completion. Each participant contributes a different fact. None should assume "
        "that forwarding a ticket transfers accountability."
    )


def add_chapter_6(doc: Document) -> None:
    doc.add_heading("6. The modern threat environment", level=1)
    add_section_intro(
        doc,
        "Digital dependence, concentrated suppliers and automation have increased both opportunity and consequence."
    )
    doc.add_heading("6.1 Dependence and attack surface", level=2)
    add_para(
        doc,
        "The attack surface of a modern organisation includes devices, web applications, APIs, identity providers, "
        "cloud platforms, remote-access paths, mobile applications, connected products, operational technology and "
        "suppliers. A small organisation may own little infrastructure yet depend on dozens of software services. "
        "This can improve security when specialist providers operate stronger controls, but it also concentrates "
        "risk: one identity, software or supply-chain failure may affect many customers."
    )
    add_para(
        doc,
        "ENISA's Threat Landscape 2025 analysed 4,875 incidents observed between July 2024 and June 2025. The "
        "dataset should not be read as a complete census, but it demonstrates the diversity of threats affecting "
        "availability, data, public administration, transport, manufacturing and other sectors. The NCSC reported "
        "429 incidents handled in 2024-25, including 204 classified as nationally significant and 18 as highly "
        "significant. The nationally significant total was more than double the previous year's 89."
    )

    doc.add_heading("6.2 Threat actors and motives", level=2)
    add_para(
        doc,
        "Threat actors include financially motivated criminals, hostile states, ideological groups, insiders, "
        "competitors and opportunistic individuals. Their capability and patience differ. Criminal groups seek "
        "payment, saleable data or fraudulent transactions. States may pursue intelligence, disruption or "
        "strategic access. Insiders may act maliciously, under coercion or through error. The same vulnerability "
        "can therefore create different risk depending on who can reach it and what outcome they seek."
    )
    add_para(
        doc,
        "Attribution is uncertain and rarely necessary for the first defensive decision. An organisation must "
        "contain malicious access before it knows the actor's identity. Threat intelligence is valuable when it "
        "changes a decision about exposure, detection or response; collections of indicators without context can "
        "consume attention without reducing risk."
    )

    doc.add_heading("6.3 Common paths to harm", level=2)
    add_para(
        doc,
        "Phishing and social engineering exploit trust, urgency and familiar communication. Credential theft "
        "turns legitimate access mechanisms against their owners. Exploitation of internet-facing vulnerabilities "
        "targets unpatched or misconfigured services. Malware and ransomware automate disruption and extortion. "
        "Supply-chain compromise reaches targets through trusted software or service providers. Cloud incidents "
        "often involve identity, permissions or exposed data rather than a failure of the underlying provider."
    )
    add_para(
        doc,
        "Connected products add long support lifecycles, weak update mechanisms, default credentials and physical "
        "effects. Operational technology can prioritise availability and safety over rapid change, making ordinary "
        "IT patching assumptions unsuitable. Mobile and home working blur the network boundary. These examples "
        "support an identity- and asset-centred model: organisations must know what they depend on and continuously "
        "verify access rather than assume that location alone creates trust."
    )

    doc.add_heading("6.4 Consequence, not novelty", level=2)
    add_para(
        doc,
        "Security discussion is easily distracted by novel tools and dramatic attack names. Many serious incidents "
        "still use familiar weaknesses: missing multi-factor authentication, excessive privilege, unpatched "
        "systems, weak recovery, exposed secrets or delayed detection. Novel capability matters, especially where "
        "AI improves scale, but consequence depends on ordinary architecture and preparedness. A convincing message "
        "becomes catastrophic only when one credential can reach critical assets and the organisation cannot "
        "detect, contain or recover."
    )


def add_chapter_7(doc: Document) -> None:
    doc.add_heading("7. Does cybersecurity work?", level=1)
    add_section_intro(
        doc,
        "Effectiveness must be assessed through avoided harm, controlled consequence and quality of recovery—not the mere existence of attacks."
    )
    doc.add_heading("7.1 What current UK evidence shows", level=2)
    add_para(
        doc,
        "The Cyber Security Breaches Survey 2025/2026 estimates that 43 per cent of UK businesses and 28 per cent "
        "of charities identified a breach or attack in the previous twelve months. Prevalence was higher among "
        "medium and large businesses. Phishing remained the most common type identified. The survey also reported "
        "increases in the proportions of affected businesses associating their most disruptive incident with "
        "revenue or share-value loss and reputational damage, although the absolute percentages remained small."
    )
    add_para(
        doc,
        "Preparedness was uneven. Fifty-two per cent of businesses had undertaken at least one listed activity to "
        "identify cyber risks, but only 30 per cent reported a cyber-risk assessment, 18 per cent a vulnerability "
        "audit and 13 per cent penetration testing. Thirty-one per cent assigned explicit board responsibility and "
        "25 per cent had a formal incident-response plan. Forty-five per cent reported none of the listed response "
        "measures. Appropriate action varies by risk, but the gap indicates that many organisations depend on "
        "digital systems without a structured understanding of failure."
    )
    add_labelled_callout(
        doc,
        "Evidence summary",
        "High attack prevalence demonstrates exposure. Low adoption of risk assessment and response planning "
        "demonstrates an implementation gap. Neither finding shows that established security practice is incapable "
        "of reducing harm.",
    )

    doc.add_heading("7.2 Why breaches do not equal total failure", level=2)
    add_para(
        doc,
        "An attempted phishing message counted as an identified attack may be blocked, reported or contained "
        "without material harm. Conversely, an organisation that reports no attacks may lack visibility. Success "
        "often appears as an event that did not escalate: a stolen password rejected by multi-factor "
        "authentication, malicious code isolated by endpoint controls, unauthorised privilege detected in logs or "
        "data restored from protected backups. These counterfactual outcomes are difficult to aggregate but central "
        "to assessing value."
    )
    add_para(
        doc,
        "The appropriate unit of analysis is therefore not simply 'breached or not breached'. Measures should "
        "include control coverage, time to detect, time to contain, scope of access, data affected, service "
        "interruption, recovery against objectives, recurrence and the quality of learning. Metrics can themselves "
        "distort behaviour: a team rewarded for fewer incidents may discourage reporting, while a team rewarded for "
        "closing vulnerabilities may prioritise easy counts over important exposure."
    )

    doc.add_heading("7.3 Evidence that controls matter", level=2)
    add_para(
        doc,
        "Security controls are not equally effective in every context, but layered protection changes attacker "
        "cost and consequence. Phishing-resistant authentication reduces the value of reusable passwords. Least "
        "privilege limits what a compromised account can reach. Segmentation restricts movement. Secure updates "
        "remove known weaknesses. Logging and detection shorten the period of unnoticed access. Offline or "
        "logically protected backups make destructive extortion less powerful. Rehearsed response improves the "
        "speed and confidence of containment."
    )
    add_para(
        doc,
        "The defence-in-depth principle recognises that each control can fail. Its value depends on independence "
        "and coverage, not the number of products purchased. Five tools relying on the same identity and telemetry "
        "may share one failure. Effective layering combines prevention, detection and recovery across technical and "
        "organisational boundaries."
    )

    doc.add_heading("7.4 The implementation gap", level=2)
    add_para(
        doc,
        "The strongest criticism of contemporary cybersecurity is not absence of knowledge but inconsistent "
        "execution. Organisations may not know their assets, maintain unsupported systems, grant excessive access, "
        "retain unnecessary data, fail to test backups or leave suppliers unexamined. Controls may be deployed "
        "without owners or tuned so poorly that staff work around them. Security programmes can become compliance "
        "theatre when evidence is produced for an audit but operational behaviour does not change."
    )
    add_para(
        doc,
        "Resource inequality matters. Large organisations can employ specialists and purchase managed services; "
        "small organisations may rely on general IT support. Secure defaults, proportionate guidance and supplier "
        "responsibility are therefore essential. A model that works only when every organisation employs a large "
        "security team cannot secure an economy composed largely of smaller organisations."
    )


def add_chapter_8(doc: Document) -> None:
    doc.add_heading("8. People, usability and security culture", level=1)
    add_section_intro(
        doc,
        "People are participants in security systems. Treating them only as weaknesses hides design and organisational causes."
    )
    doc.add_heading("8.1 The problem with 'human error'", level=2)
    add_para(
        doc,
        "Post-incident explanations often end with a person clicking, sharing, misconfiguring or delaying. The "
        "description may be factually correct but analytically incomplete. Actions occur inside interfaces, "
        "workloads, incentives, authority structures and technical boundaries designed by organisations. If one "
        "ordinary mistake permits catastrophic and irreversible harm, excessive privilege, weak isolation, poor "
        "detection or absent recovery also contributed."
    )
    add_para(
        doc,
        "Blame can reduce security by discouraging early reporting. A person who fears punishment may hide an "
        "error until evidence disappears or harm spreads. A learning culture makes rapid reporting easy, limits "
        "access by default and examines why the system allowed the action. This does not remove individual "
        "responsibility for deliberate misconduct; it distinguishes accountability from a simplistic search for "
        "the last person who touched the system."
    )

    doc.add_heading("8.2 The cost of security advice", level=2)
    add_para(
        doc,
        "Herley's analysis of user security advice argues that precautions impose time and effort on users while "
        "many benefits are uncertain or accrue to others. Ignoring advice can therefore be rational from an "
        "individual perspective. The lesson is not that users should reject security, but that designers must "
        "consider the economic and cognitive cost of controls. Repeated warnings, complex password rules and "
        "frequent low-value prompts train people to dismiss interruption."
    )
    add_para(
        doc,
        "Usable security makes the safer action the easier and more comprehensible action. Password managers, "
        "single sign-on, phishing-resistant authentication, clear permission requests, automatic updates and safe "
        "recovery reduce reliance on memory and vigilance. Defaults matter because most users reasonably expect "
        "products to work without expert reconfiguration."
    )

    doc.add_heading("8.3 Education and awareness", level=2)
    add_para(
        doc,
        "Education remains necessary. People need to recognise suspicious requests, understand the value of "
        "information, know how to report concerns and appreciate why certain controls exist. Training is most "
        "effective when relevant to actual work, repeated at sensible intervals and supported by immediate "
        "feedback. Generic annual presentations cannot prepare every role for changing attacks."
    )
    add_para(
        doc,
        "Awareness is not a substitute for engineering. Training cannot repair unsupported software, remove a "
        "default password, limit a cloud permission or create a tested backup. Organisations sometimes overstate "
        "the user's duty because behaviour change appears cheaper than redesign. A balanced programme combines "
        "education with technical guardrails, clear process and leadership."
    )

    doc.add_heading("8.4 Security culture", level=2)
    add_para(
        doc,
        "Culture is visible in routine decisions: whether deadlines override risk without approval, whether staff "
        "can challenge unusual requests, whether incidents are reported quickly, whether leaders follow the same "
        "controls as everyone else and whether security teams understand operational needs. A culture of fear can "
        "produce hidden workarounds; a culture of convenience can normalise excessive access. Healthy culture "
        "supports constructive challenge and treats security as part of quality."
    )


def add_chapter_9(doc: Document) -> None:
    doc.add_heading("9. Technical architecture and secure-by-design practice", level=1)
    add_section_intro(
        doc,
        "Technology is effective when controls are designed into systems, maintained as dependencies change and supported by evidence."
    )
    doc.add_heading("9.1 Secure by design and by default", level=2)
    add_para(
        doc,
        "Secure by design moves security decisions earlier in a product or service lifecycle. Teams identify "
        "assets and abuse cases, model threats, choose trust boundaries, minimise privilege, validate input, "
        "protect secrets, log relevant events and plan updates before deployment. Secure by default means the "
        "product begins in a defensible state without requiring specialist knowledge from every customer."
    )
    add_para(
        doc,
        "NCSC software-security guidance asks organisations to establish ownership, threat-model throughout "
        "development, prevent common weaknesses, protect development environments and manage vulnerabilities. "
        "CISA's secure-by-design campaign similarly argues that customer security should be a core business "
        "requirement. This changes the distribution of responsibility: a manufacturer is better placed than each "
        "customer to remove a default weakness affecting every installation."
    )

    doc.add_heading("9.2 Identity and access", level=2)
    add_para(
        doc,
        "Identity has become a primary security boundary in cloud and hybrid environments. Strong authentication "
        "must be combined with lifecycle management, least privilege, separation of duties, review of dormant "
        "accounts and protection of administrative access. Multi-factor authentication reduces risk but methods "
        "differ; phishing-resistant approaches provide stronger protection than codes that can be relayed."
    )
    add_para(
        doc,
        "Zero-trust principles are often reduced to a product label. Their practical value lies in removing "
        "implicit trust based only on network location, verifying access in context, minimising privilege and "
        "assuming that compromise can occur. Continuous verification should not mean constant obstruction. "
        "Signals and policy must be reliable enough that users can perform legitimate work."
    )

    doc.add_heading("9.3 Vulnerability and configuration management", level=2)
    add_para(
        doc,
        "Known vulnerabilities create avoidable risk when exposed systems remain unpatched, but patching is not a "
        "single metric. Organisations need an inventory, supplier information, exposure context, tested deployment "
        "and exception handling. Configuration weaknesses—public storage, unnecessary services, excessive "
        "permissions and unmanaged secrets—can be as important as software defects. Baselines and automated checks "
        "help teams detect drift."
    )
    add_para(
        doc,
        "Prioritisation should combine technical severity with exploitation evidence, reachability, asset value and "
        "available mitigations. Treating every finding as equally urgent overwhelms teams and reduces attention to "
        "the attack paths most likely to create serious consequence."
    )

    doc.add_heading("9.4 Monitoring and detection", level=2)
    add_para(
        doc,
        "Logs provide evidence only when relevant events are generated, protected, retained and reviewed. A "
        "security operations centre may combine endpoint, identity, network, application and cloud telemetry to "
        "identify behaviour inconsistent with expected use. Detection engineering converts threat hypotheses into "
        "testable signals. Too many low-quality alerts create fatigue and hide important events."
    )
    add_para(
        doc,
        "Monitoring must also respect privacy and employment law. Collecting every possible event can increase "
        "risk, cost and intrusion without improving decisions. Purpose, access, retention and oversight should be "
        "defined. The aim is sufficient evidence for prevention, investigation and learning—not surveillance as "
        "an end in itself."
    )

    doc.add_heading("9.5 Response, continuity and recovery", level=2)
    add_para(
        doc,
        "Incident response requires prepared roles, communication routes, technical access and decision authority. "
        "Plans should cover triage, containment, evidence, legal assessment, supplier contact, affected people, "
        "service restoration and public communication. Exercises expose missing permissions and assumptions before "
        "a crisis. A plan that has never been tested remains a hypothesis."
    )
    add_para(
        doc,
        "Backups are useful only when recoverable, sufficiently recent and protected from the same compromise as "
        "the primary system. Recovery objectives should reflect operational need. Restoring service without "
        "understanding the cause may reintroduce an attacker; delaying restoration indefinitely can create greater "
        "harm. Resilience balances trusted recovery with continuity of essential functions."
    )

    doc.add_heading("9.6 Connected-product security", level=2)
    add_para(
        doc,
        "Consumer connected products historically placed substantial responsibility on purchasers who could not "
        "evaluate update policy, vulnerability handling or hidden dependencies. The UK's Product Security and "
        "Telecommunications Infrastructure product-security regime, effective from 29 April 2024, introduced "
        "minimum obligations for relevant consumer connectable products, including requirements associated with "
        "passwords, vulnerability reporting and transparency about security updates."
    )
    add_para(
        doc,
        "Regulation does not make every product secure, but it changes the baseline and creates accountable "
        "obligations upstream. Procurement and provisioning still matter: organisations should understand support "
        "periods, network requirements, data flows, ownership, update mechanisms and secure configuration before "
        "deployment. This is where security research can become a practical decision tool rather than general "
        "advice."
    )


def add_chapter_10(doc: Document) -> None:
    doc.add_heading("10. Automation and artificial intelligence", level=1)
    add_section_intro(
        doc,
        "Automation is required for scale, but it inherits assumptions and cannot own consequence."
    )
    doc.add_heading("10.1 Defensive automation", level=2)
    add_para(
        doc,
        "Modern estates produce more events, assets and configuration changes than people can inspect manually. "
        "Automation supports asset discovery, software inventory, configuration checking, vulnerability scanning, "
        "log correlation, malware analysis, alert enrichment and repeatable response. It can reduce delay and "
        "variation, particularly for high-volume tasks with clear rules."
    )
    add_para(
        doc,
        "Automation also creates failure modes. An incorrect rule can block legitimate service, remove evidence or "
        "disable accounts at scale. A detector trained on past behaviour may miss a novel attack or reproduce bias "
        "in its data. Safe automation therefore uses bounded authority, testing, audit trails, rollback and human "
        "review proportionate to consequence."
    )

    doc.add_heading("10.2 AI for defenders and attackers", level=2)
    add_para(
        doc,
        "Artificial intelligence can assist defenders with summarisation, code review, anomaly identification, "
        "triage, query generation and threat-intelligence analysis. Generative systems can make specialist tools "
        "more accessible by translating between natural language and technical formats. They may also help smaller "
        "teams interpret evidence more quickly."
    )
    add_para(
        doc,
        "The same capabilities can help attackers create persuasive content, translate lures, search public "
        "information, generate code and automate reconnaissance. AI does not eliminate the need for access, "
        "execution and persistence, but it can reduce cost and increase scale. Defensive planning should focus on "
        "the resulting attack path rather than assume that every AI-generated artefact is uniquely detectable."
    )

    doc.add_heading("10.3 Securing AI systems", level=2)
    add_para(
        doc,
        "AI systems introduce assets and dependencies of their own: training and retrieval data, model artefacts, "
        "prompts, tools, credentials, external APIs, user inputs and generated outputs. Risks include data leakage, "
        "poisoning, prompt injection, unsafe tool use, unreliable output and over-trust. Traditional controls such "
        "as access management, input handling, logging, change control and incident response remain relevant but "
        "must be adapted to probabilistic behaviour."
    )
    add_para(
        doc,
        "NIST's AI Risk Management Framework organises work around Govern, Map, Measure and Manage. Its value is "
        "the insistence that risk be considered across the lifecycle and in context. An accurate model can still "
        "be unsafe if used for a decision it was not designed to support. Accountability must remain with people "
        "and organisations able to explain the purpose, evidence and consequence of deployment."
    )

    doc.add_heading("10.4 Will AI replace security professionals?", level=2)
    add_para(
        doc,
        "Some tasks will change or become automated, particularly repetitive analysis and first-pass production. "
        "However, security work includes uncertainty, negotiation, ethics, system context and responsibility for "
        "consequence. A model cannot accept organisational risk or repair a broken relationship between product, "
        "IT and leadership. The more plausible future is collaboration: professionals use AI to extend reach while "
        "remaining accountable for validation and decisions."
    )


def add_chapter_11(doc: Document) -> None:
    doc.add_heading("11. Governance, law and ethical limits", level=1)
    add_section_intro(
        doc,
        "Security becomes effective when technical evidence informs accountable organisational decisions."
    )
    doc.add_heading("11.1 Governance and risk ownership", level=2)
    add_para(
        doc,
        "Cyber risk can affect revenue, safety, legal duties, reputation and the ability to operate. It is therefore "
        "an enterprise risk rather than a matter that leadership can delegate completely to an IT team. The UK "
        "Cyber Governance Code of Practice sets expectations for boards and directors across risk management, "
        "strategy, people, incident planning and assurance. Its emphasis reflects a basic principle: specialists "
        "advise and operate controls, but organisational leaders decide priorities and accept residual risk."
    )
    add_para(
        doc,
        "Governance should connect risk appetite to technical work. A board-level statement that disruption is "
        "intolerable must become recovery objectives, architecture, investment, supplier requirements and "
        "exercises. Reporting should describe exposure, control confidence and consequence rather than present "
        "uninterpreted counts. Leadership also needs uncertainty; a falsely precise dashboard can hide weak data."
    )

    doc.add_heading("11.2 Law, regulation and standards", level=2)
    add_para(
        doc,
        "Law establishes minimum duties and remedies, while standards organise good practice. Data-protection law "
        "requires appropriate security for personal data but also principles such as purpose limitation and data "
        "minimisation. Product-security regulation places obligations on relevant manufacturers and supply-chain "
        "businesses. Sector rules may add resilience, reporting or safety duties. Compliance can motivate action, "
        "but a control designed only to pass an audit may fail against real threats."
    )
    add_para(
        doc,
        "Frameworks such as NIST CSF 2.0 are intentionally adaptable. They describe outcomes rather than prescribe "
        "one product set. This supports proportionality but requires judgement. Certification and assurance can "
        "increase confidence, yet neither proves that a system is invulnerable or that controls remain effective "
        "after change."
    )

    doc.add_heading("11.3 Ethical tensions", level=2)
    add_para(
        doc,
        "Security decisions distribute power. Monitoring can reveal misconduct but also expose personal behaviour. "
        "Access controls protect information but may exclude people during urgent need. Automated fraud controls "
        "reduce loss but can unfairly deny legitimate users. Vulnerability research can improve products but create "
        "risk if disclosure is mishandled. Ethical practice asks who benefits, who bears cost, whose rights are "
        "affected and what remedy exists when a control is wrong."
    )
    add_para(
        doc,
        "Proportionality provides a useful test. Data collection and restriction should be necessary for a defined "
        "purpose, no more intrusive than required, time-limited where appropriate and subject to oversight. Security "
        "is not an automatic justification for unlimited surveillance or secrecy. Trustworthy protection includes "
        "the ability to challenge decisions and hold powerful actors accountable."
    )

    doc.add_heading("11.4 Inequality and public value", level=2)
    add_para(
        doc,
        "Security burdens fall unevenly. People with limited digital access, time or technical confidence may be "
        "least able to navigate complex recovery. Small organisations may lack specialist staff. Victims of abuse "
        "face threats from people who know their devices and relationships. Product design and public policy should "
        "therefore avoid assuming an expert, well-resourced user. Recovery routes, accessible authentication and "
        "safe support are part of effective security."
    )


def add_chapter_12(doc: Document) -> None:
    doc.add_heading("12. Integrated incident analysis", level=1)
    add_section_intro(
        doc,
        "A stolen cloud credential demonstrates why the four fields must retain distinct questions while operating as one response."
    )
    doc.add_heading("12.1 Scenario", level=2)
    add_para(
        doc,
        "An employee receives a convincing message that imitates a supplier's document-sharing notification. The "
        "employee follows the link and completes an authentication flow controlled by an attacker. The attacker "
        "uses the resulting session to enter the organisation's cloud environment, searches email and shared "
        "storage, creates a persistence mechanism and downloads files containing customer and staff information. "
        "Unusual activity is detected several hours later."
    )
    add_para(
        doc,
        "The scenario is deliberately ordinary. It does not depend on an unknown vulnerability or extraordinary "
        "actor. Its consequence depends on the authentication method, device evidence, account privilege, cloud "
        "configuration, information holdings, detection coverage, supplier process and response authority."
    )

    doc.add_heading("12.2 Computer-security perspective", level=2)
    add_para(
        doc,
        "Computer security asks whether the endpoint was compromised, what executed locally, whether browser or "
        "credential material remains exposed and whether the system can be trusted. Endpoint logs, process history, "
        "browser artefacts, malware scans and memory or disk evidence may be relevant. The endpoint team isolates "
        "the device if necessary and determines whether rebuild is safer than continued use."
    )

    doc.add_heading("12.3 IT-security perspective", level=2)
    add_para(
        doc,
        "IT security asks how the identity was authenticated, which permissions it held, what services were reached "
        "and whether persistence or forwarding rules remain. Identity administrators revoke sessions, reset "
        "credentials, review multi-factor methods, search related sign-ins and remove malicious changes. Cloud and "
        "network teams preserve logs, restrict access and examine whether least privilege and conditional policy "
        "worked as intended."
    )

    doc.add_heading("12.4 Information-security perspective", level=2)
    add_para(
        doc,
        "Information security identifies the files accessed, their classification, owners, retention basis and "
        "contractual or legal significance. It assesses confidentiality, integrity and availability impact, "
        "coordinates with data-protection and legal specialists, and records risk decisions. The analysis may "
        "reveal that excessive data retention increased consequence even though the technical entry path was an "
        "identity failure."
    )

    doc.add_heading("12.5 Cybersecurity perspective", level=2)
    add_para(
        doc,
        "Cybersecurity connects the event into an adversarial incident. Responders construct a timeline, search for "
        "related activity, identify tactics, coordinate containment, assess wider campaigns and communicate with "
        "suppliers or authorities where appropriate. Threat intelligence may establish whether the lure belongs to "
        "a known campaign, but response should not wait for confident attribution."
    )

    doc.add_heading("12.6 Governance, people and recovery", level=2)
    add_para(
        doc,
        "The incident lead needs authority to make cross-system decisions. Management must balance containment with "
        "service continuity and decide notifications. Communications should support affected people without "
        "speculation. The employee should be treated as an evidence source and participant, not the complete cause. "
        "Review should ask why the authentication could be relayed, why the account could reach the data, why "
        "detection took hours and whether the retained information was necessary."
    )
    add_para(
        doc,
        "A mature outcome includes revoked access, trusted devices, validated logs, notification decisions, restored "
        "service and tracked improvements. Possible actions include phishing-resistant authentication, tighter "
        "permissions, safer supplier workflows, improved detection, reduced retention and a rehearsed cloud "
        "containment process. The incident crosses every field, but the distinct lenses prevent important questions "
        "from disappearing."
    )

    doc.add_heading("12.7 What the scenario proves", level=2)
    add_para(
        doc,
        "No single control guarantees prevention. Effectiveness emerges from a system of barriers and recovery. If "
        "the message is blocked, harm is avoided. If the employee reports quickly, response begins earlier. If the "
        "credential cannot satisfy phishing-resistant authentication, access fails. If privilege is limited, "
        "consequence shrinks. If monitoring detects unusual activity, dwell time falls. If information is minimised, "
        "fewer people are affected. Security is therefore cumulative and socio-technical."
    )


def add_chapter_13(doc: Document) -> None:
    doc.add_heading("13. Discussion", level=1)
    add_section_intro(
        doc,
        "The evidence supports a conditional answer and clarifies where contemporary practice succeeds or fails."
    )
    doc.add_heading("13.1 Is cybersecurity secure?", level=2)
    add_para(
        doc,
        "Cybersecurity is not itself an object that can be certified as permanently secure. It is a discipline for "
        "managing digital risk. Asking whether it is 'secure' confuses method with outcome. The stronger question is "
        "whether the discipline offers practices capable of reducing unacceptable harm in real environments. The "
        "evidence supports that capability. Authentication, least privilege, secure design, vulnerability "
        "management, monitoring, response and recovery each address known paths to harm, while governance connects "
        "them to organisational purpose."
    )
    add_para(
        doc,
        "However, capability does not guarantee implementation. Current UK evidence shows that many organisations "
        "lack structured risk assessment and formal response planning. Threat activity and nationally significant "
        "incidents remain serious. Cybersecurity therefore provides effective security only where controls are "
        "selected for context, maintained, tested and supported by accountable leadership."
    )

    doc.add_heading("13.2 What the field distinctions contribute", level=2)
    add_para(
        doc,
        "Separating computer, IT, information and cybersecurity improves analysis. Computer security prevents the "
        "individual system from becoming a black box. IT security addresses the estate and operational service. "
        "Information security protects meaning, obligations and lifecycle. Cybersecurity connects adversarial "
        "digital activity to wider harm. The distinctions are not competing definitions; they are complementary "
        "starting points."
    )
    add_para(
        doc,
        "The model also exposes gaps. A vulnerability can be technically fixed while unnecessary information "
        "remains. A policy can be compliant while devices are unmanaged. A SOC can detect activity but lack "
        "authority to isolate a business service. Clear roles and handoffs ensure that evidence reaches the person "
        "able to decide."
    )

    doc.add_heading("13.3 People and technology", level=2)
    add_para(
        doc,
        "The investigation rejects both technological determinism and user blame. Tools change exposure but inherit "
        "design assumptions. People make mistakes but also detect anomalies, report context and exercise ethical "
        "judgement. Automation is necessary for scale, while accountable human oversight is necessary for "
        "consequence. Effective security designs for this partnership."
    )

    doc.add_heading("13.4 Prevention and resilience", level=2)
    add_para(
        doc,
        "The most important change in the understanding of security is the move from perfect prevention to "
        "resilience. Prevention remains valuable; accepting that incidents can occur is not permission for weak "
        "controls. It is recognition that complex systems need detection, containment and trusted recovery. An "
        "organisation able to restore essential service and learn from failure is more secure than one whose only "
        "claim is that no incident has yet been observed."
    )

    doc.add_heading("13.5 Product and policy implications", level=2)
    add_para(
        doc,
        "Responsibility should move towards actors able to reduce risk efficiently. Manufacturers should provide "
        "secure defaults and updates. Service providers should expose usable evidence. Organisations should make "
        "ownership explicit and test response. Government should establish proportionate baselines and support "
        "smaller organisations. Users should receive understandable controls and recovery, not an endless transfer "
        "of responsibility."
    )
    add_para(
        doc,
        "This principle also explains the practical direction of device-comparison and provisioning tools. Advice "
        "becomes valuable when it helps a person compare support, understand risk, configure securely and record a "
        "decision. The aim is not to certify perfect safety but to turn evidence into the next safer action."
    )

    doc.add_heading("13.6 Further research", level=2)
    add_para(
        doc,
        "Future research should test the four-field model with practitioners and smaller organisations, examine "
        "whether clearer handoffs improve incident outcomes, and evaluate how security evidence can be presented "
        "to non-specialist decision makers. Longitudinal work is needed to determine whether product-security "
        "regulation changes update support and vulnerability handling in practice. AI-assisted security tools also "
        "require empirical evaluation of accuracy, operator trust, auditability and failure."
    )


def add_chapter_14(doc: Document) -> None:
    doc.add_heading("14. Conclusion", level=1)
    add_para(
        doc,
        "This dissertation asked whether contemporary cybersecurity can provide effective security for modern "
        "digital life and how four overlapping fields combine to produce that protection. The answer begins with "
        "a correction: security is not a permanent guarantee and cybersecurity is not an object that can be labelled "
        "secure or insecure. Security is the justified, continuously reviewed reduction of unacceptable harm. It "
        "requires prevention, detection, containment and recovery."
    )
    add_para(
        doc,
        "Computer security protects the individual system; IT security protects the organisational technology "
        "estate; information security protects information and obligations across all formats; cybersecurity "
        "addresses harm through interconnected digital systems and its wider consequences. Their boundaries overlap "
        "because incidents cross them. The value of the distinctions lies in clearer questions, expertise, evidence "
        "and accountability."
    )
    add_para(
        doc,
        "Current evidence demonstrates serious pressure. Breaches and attacks are widely identified, nationally "
        "significant incidents have increased and many organisations lack risk assessment or formal response plans. "
        "Yet these facts do not show that cybersecurity is inherently ineffective. Established controls change "
        "attacker cost, reduce privilege, expose malicious activity and support recovery. The central weakness is "
        "uneven adoption and maintenance."
    )
    add_para(
        doc,
        "Technology alone is insufficient. Security depends on design, usable defaults, informed people, responsible "
        "leadership, lawful processing, supplier behaviour and tested recovery. AI and automation can extend "
        "defensive capacity but do not own risk or consequence. Education helps people participate, but must not be "
        "used to excuse unsafe products or fragile architecture."
    )
    add_para(
        doc,
        "The final answer is therefore conditional but positive. Contemporary cybersecurity can provide effective "
        "security when it is treated as a continuous socio-technical capability: governed as risk, engineered into "
        "systems, measured through evidence, designed for real users and prepared for failure. It cannot keep pace "
        "when reduced to a purchased tool, an annual exercise or a demand that individuals never make mistakes. The "
        "most credible security does not promise that nothing will happen. It demonstrates that avoidable harm is "
        "made less likely, important change is discovered, failure is contained and trusted operation can return."
    )


def add_references(doc: Document) -> None:
    doc.add_heading("References", level=1)
    references = [
        ("Adams, A. and Sasse, M.A. (1999).", "Users are not the enemy. Communications of the ACM, 42(12), 40-46.", "https://doi.org/10.1145/322796.322806"),
        ("CISA (2023).", "Shifting the Balance of Cybersecurity Risk: Principles and Approaches for Security-by-Design and -Default.", "https://www.cisa.gov/resources-tools/resources/secure-by-design"),
        ("CISA and FBI (2025).", "Product Security Bad Practices: updated guidance.", "https://www.cisa.gov/news-events/alerts/2025/01/17/cisa-and-fbi-release-updated-guidance-product-security-bad-practices"),
        ("Department for Science, Innovation and Technology (DSIT) (2024).", "The UK Product Security and Telecommunications Infrastructure product-security regime.", "https://www.gov.uk/government/publications/the-uk-product-security-and-telecommunications-infrastructure-product-security-regime"),
        ("DSIT and Home Office (2026).", "Cyber Security Breaches Survey 2025/2026.", "https://www.gov.uk/government/statistics/cyber-security-breaches-survey-20252026/cyber-security-breaches-survey-20252026"),
        ("DSIT and NCSC (2025).", "Cyber Governance Code of Practice.", "https://www.gov.uk/government/publications/cyber-governance-code-of-practice"),
        ("European Union Agency for Cybersecurity (ENISA) (2025).", "ENISA Threat Landscape 2025.", "https://www.enisa.europa.eu/publications/enisa-threat-landscape-2025"),
        ("Herley, C. (2009).", "So Long, and No Thanks for the Externalities: The Rational Rejection of Security Advice by Users.", "https://www.microsoft.com/en-us/research/publication/so-long-and-no-thanks-for-the-externalities-the-rational-rejection-of-security-advice-by-users/"),
        ("Kissel, R. et al. (2013).", "NISTIR 7298 Revision 2: Glossary of Key Information Security Terms.", "https://doi.org/10.6028/NIST.IR.7298r2"),
        ("NCSC (2018).", "Secure by Default.", "https://www.ncsc.gov.uk/information/secure-default"),
        ("NCSC (2023).", "Cyber security design principles.", "https://www.ncsc.gov.uk/collection/cyber-security-design-principles"),
        ("NCSC (2025a).", "Annual Review 2025: Incident management.", "https://www.ncsc.gov.uk/collection/ncsc-annual-review-2025/chapter-01-cyber-threat-to-the-uk/incident-management"),
        ("NCSC (2025b).", "Software Security Code of Practice: secure design and development.", "https://www.ncsc.gov.uk/collection/software-security-code-of-practice-implementation-guidance/theme-1-secure-design-development"),
        ("NIST (2020).", "SP 800-207: Zero Trust Architecture.", "https://doi.org/10.6028/NIST.SP.800-207"),
        ("NIST (2023).", "Artificial Intelligence Risk Management Framework (AI RMF 1.0).", "https://doi.org/10.6028/NIST.AI.100-1"),
        ("NIST (2024a).", "The NIST Cybersecurity Framework (CSF) 2.0.", "https://doi.org/10.6028/NIST.CSWP.29"),
        ("NIST (2024b).", "Artificial Intelligence Risk Management Framework: Generative Artificial Intelligence Profile.", "https://doi.org/10.6028/NIST.AI.600-1"),
        ("NIST (2026).", "NICE Framework Components, version 2.2.0.", "https://www.nist.gov/itl/applied-cybersecurity/nice/nice-framework-resource-center/nice-framework-current-versions"),
        ("Petersen, R. et al. (2020).", "NIST SP 800-181 Rev. 1: Workforce Framework for Cybersecurity (NICE Framework).", "https://doi.org/10.6028/NIST.SP.800-181r1"),
        ("Sasse, M.A., Brostoff, S. and Weirich, D. (2001).", "Transforming the 'weakest link': a human/computer interaction approach to usable and effective security.", "https://doi.org/10.1023/A:1011902718709"),
        ("UK Parliament (2018).", "Data Protection Act 2018.", "https://www.legislation.gov.uk/ukpga/2018/12/contents"),
        ("UK Parliament (2022).", "Product Security and Telecommunications Infrastructure Act 2022.", "https://www.legislation.gov.uk/ukpga/2022/46/contents"),
        ("Verizon (2025).", "2025 Data Breach Investigations Report.", "https://www.verizon.com/business/resources/reports/dbir/"),
        ("von Solms, R. and van Niekerk, J. (2013).", "From information security to cyber security. Computers & Security, 38, 97-102.", "https://doi.org/10.1016/j.cose.2013.04.004"),
        ("World Economic Forum (2025).", "Global Cybersecurity Outlook 2025.", "https://www.weforum.org/publications/global-cybersecurity-outlook-2025/"),
    ]
    for reference in references:
        add_reference(doc, *reference)
    note = doc.add_paragraph(
        "All web sources were last checked on 19 July 2026. Statistical figures should be refreshed before any later formal submission."
    )
    note.style = doc.styles["Source note"]


def add_appendices(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix A. Field comparison", level=1)
    for label, body in (
        ("Computer security", "Protected object: an individual computing system and locally processed data. Core evidence: trusted build, hardened configuration, endpoint telemetry, vulnerability and malware analysis. Representative roles: endpoint or platform security engineer, operating-system specialist, malware analyst, vulnerability researcher."),
        ("IT security", "Protected object: an organisation's technology estate and operational services. Core evidence: inventory, identity records, configuration baselines, monitoring, access reviews, patch and recovery tests. Representative roles: identity specialist, network or cloud security engineer, security administrator, SOC analyst, vulnerability manager."),
        ("Information security", "Protected object: information in any form and the obligations attached to it. Core evidence: classification, risk assessment, policy, assurance, supplier review, retention and continuity decisions. Representative roles: information-security manager, risk analyst, security architect, auditor, information owner."),
        ("Cybersecurity", "Protected object: people, organisations, services and infrastructure exposed to harm through interconnected digital systems. Core evidence: threat assessment, incident timeline, attack-path findings, containment and lessons learned. Representative roles: incident responder, threat hunter, intelligence analyst, penetration tester, forensics specialist, OT cybersecurity engineer."),
    ):
        add_labelled_callout(doc, label, body)

    doc.add_heading("Appendix B. Security effectiveness test", level=1)
    add_restarted_numbered(
        doc,
        (
            "Purpose: What valued asset, service, person or right is being protected?",
            "Threat: Which plausible actor, event or failure could cause unacceptable harm?",
            "Ownership: Who has authority to decide and who operates the relevant control?",
            "Prevention: Which controls reduce likelihood and how is their coverage demonstrated?",
            "Detection: Which evidence would reveal harmful change and who reviews it?",
            "Containment: What limits the scope of a compromised device, identity or supplier?",
            "Response: Are roles, permissions, communications and legal routes rehearsed?",
            "Recovery: Can trusted operation and necessary information be restored within agreed objectives?",
            "Usability and ethics: Can people use the control, and is it lawful, necessary and proportionate?",
            "Learning: Are incidents, exercises and control failures converted into owned improvements?",
        ),
    )

    doc.add_heading("Appendix C. Relationship to the original project", level=1)
    add_para(
        doc,
        "The 2020 A-level EPQ established the enduring starting question: why does cybersecurity appear unable to "
        "provide complete protection, and what part do people, education and automation play? The present "
        "dissertation is not an edited copy or annotated commentary. It replaces the original structure, removes "
        "the unsupported questionnaire claims, uses a new research method, introduces explicit field boundaries "
        "and evaluates evidence available in 2026. The original submission remains archived so the development of "
        "the research can be examined transparently."
    )


def build() -> Path:
    doc = Document()
    title = "Can contemporary cybersecurity provide effective security for modern digital life?"
    configure_document(doc, title, "Full updated Extended Project Qualification dissertation")
    doc.core_properties.keywords = (
        "EPQ, cybersecurity, computer security, IT security, information security, risk, resilience, security roles"
    )

    # Narrative-proposal token map adapted to A4 academic publication.
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"].font.size = Pt(10.2)
    doc.styles["Normal"].paragraph_format.space_after = Pt(6)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.18
    doc.styles["Normal"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.styles["Heading 1"].font.size = Pt(18)
    doc.styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.styles["Heading 1"].paragraph_format.space_before = Pt(17)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(8)
    doc.styles["Heading 2"].font.size = Pt(12.5)
    doc.styles["Heading 2"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(11)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(5)
    doc.styles["Heading 3"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    header = section.header.paragraphs[0]
    header.text = "CAN CONTEMPORARY CYBERSECURITY PROVIDE EFFECTIVE SECURITY?"
    header.style = doc.styles["Source note"]
    footer = section.footer.paragraphs[0]
    footer.runs[0].text = "NATHAN BROWN-BENNETT  /  2026"

    add_cover(doc)
    add_front_matter(doc)

    body = doc.add_section(WD_SECTION.NEW_PAGE)
    body.top_margin = Cm(1.65)
    body.bottom_margin = Cm(1.6)
    body.left_margin = Cm(1.7)
    body.right_margin = Cm(1.7)
    body.header_distance = Cm(0.75)
    body.footer_distance = Cm(0.75)
    set_section_columns(body, 2, 460)

    for chapter in (
        add_chapter_1,
        add_chapter_2,
        add_chapter_3,
        add_chapter_4,
        add_chapter_5,
        add_chapter_6,
        add_chapter_7,
        add_chapter_8,
        add_chapter_9,
        add_chapter_10,
        add_chapter_11,
        add_chapter_12,
        add_chapter_13,
        add_chapter_14,
        add_references,
        add_appendices,
    ):
        chapter(doc)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
